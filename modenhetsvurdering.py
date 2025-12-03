"""
MODENHETSVURDERING - GEVINSTREALISERING
Gjennomføres i samarbeid med konsern økonomi og digital transformasjon
"""

import subprocess
import sys

# Automatisk installasjon av nødvendige pakker
def install_package(package_name):
    try:
        subprocess.check_call([sys.executable, "-m", "pip", "install", package_name, "-q"])
        return True
    except:
        return False

# Sjekk og installer python-docx
try:
    from docx import Document
    DOCX_AVAILABLE = True
except ImportError:
    print("Installerer python-docx...")
    if install_package("python-docx"):
        try:
            from docx import Document
            DOCX_AVAILABLE = True
        except:
            DOCX_AVAILABLE = False
    else:
        DOCX_AVAILABLE = False

# Sjekk og installer reportlab
try:
    from reportlab.lib.pagesizes import A4
    REPORTLAB_AVAILABLE = True
except ImportError:
    print("Installerer reportlab...")
    if install_package("reportlab"):
        try:
            from reportlab.lib.pagesizes import A4
            REPORTLAB_AVAILABLE = True
        except:
            REPORTLAB_AVAILABLE = False
    else:
        REPORTLAB_AVAILABLE = False

import streamlit as st
import pandas as pd
import plotly.graph_objects as go
import numpy as np
from datetime import datetime
import pickle
import os
import tempfile
import base64
from io import BytesIO

# ============================================================================
# KONFIGURASJON
# ============================================================================
st.set_page_config(
    page_title="Modenhetsvurdering - Bane NOR",
    page_icon="🚂",
    layout="wide",
    initial_sidebar_state="expanded"
)

DATA_FILE = "modenhet_data_v5.pkl"

# ============================================================================
# BANE NOR FARGEPALETT
# ============================================================================
COLORS = {
    'primary_dark': '#172141',
    'primary': '#0053A6',
    'primary_light': '#64C8FA',
    'success': '#35DE6D',
    'warning': '#FFA040',
    'danger': '#FF6B6B',
    'white': '#FFFFFF',
    'black': '#000000',
    'gray_light': '#F2FAFD',
    'gray': '#E8E8E8'
}

# ============================================================================
# BANE NOR IKONER (SVG)
# ============================================================================
ICONS = {
    # Tog-ikon for header
    'train': '''<svg viewBox="0 0 141.73 141.73" xmlns="http://www.w3.org/2000/svg"><path d="M98.51 77.81 44.51 77.81 44.51 70.6 85.51 70.6 85.51 68.08 44.51 68.14 44.51 52 30.85 52 30.85 96.33 44.49 96.33 44.49 80.27 98.49 80.27 98.49 96.33 112.12 96.33 112.12 52 98.51 52ZM42 93.82 33.4 93.82 33.4 54.54 42 54.54ZM101 54.54 109.61 54.54 109.61 93.82 101 93.82ZM22.08 65.3 24.54 65.3 24.54 83.06 22.08 83.06ZM120.91 65.3 120.91 83.06 118.45 83.06 118.45 65.3Z" fill="{color}"/></svg>''',
    
    # Planlegging - Dashboard/skjerm
    'planning': '''<svg viewBox="0 0 141.73 141.73" xmlns="http://www.w3.org/2000/svg"><path d="M82.75 36.66 82.75 105.2 121.46 105.2 121.46 36.66ZM118.96 102.66 85.25 102.66 85.25 90.07 119 90.07ZM118.96 87.53 85.25 87.53 85.25 39.16 119 39.16ZM106.63 44.26 97.63 44.26 97.63 41.76 106.63 41.76ZM102.1 101.26C104.652 101.26 106.72 99.19 106.72 96.64 106.72 94.09 104.65 92.02 102.1 92.02 99.55 92.02 97.48 94.09 97.48 96.64 97.5 99.19 99.56 101.23 102.1 101.23ZM102.1 94.53C103.27 94.52 104.23 95.46 104.24 96.63 104.25 97.8 103.31 98.76 102.14 98.77 100.97 98.78 100.01 97.84 100 96.67L100 96.62C100 95.46 100.94 94.51 102.1 94.5ZM21.58 97.72 47.29 97.72 47.29 102.56 41.64 102.56 41.64 105.06 55.44 105.06 55.44 102.56 49.79 102.56 49.79 97.72 75.49 97.72 75.49 63 21.58 63ZM24.08 65.5 73 65.5 73 95.22 24.08 95.22Z" fill="{color}"/></svg>''',
    
    # Gjennomføring - Bygninger/by
    'execution': '''<svg viewBox="0 0 141.73 141.73" xmlns="http://www.w3.org/2000/svg"><path d="M120.45 106.38 120.45 87 113.15 45 81.76 45 84.08 59 67.7 59 65.33 44.87 21.1 44.94 26.44 76.94 26.44 108.87 120.44 108.87 120.44 106.37ZM24.06 47.44 63.22 47.37 69.89 87 56.89 87 56.89 75.6 28.73 75.6ZM54.38 87 54.38 106.43 42.91 106.43 42.91 87 40.41 87 40.41 106.43 28.94 106.43 28.94 78.1 54.38 78.1ZM85.6 106.38 72.6 106.38 72.6 96.67 70.1 96.67 70.1 106.38 56.88 106.38 56.88 89.45 85.65 89.45ZM85.65 87 72.42 87 68.12 61.47 84.49 61.47 85.71 68.78ZM88 67.38 84.71 47.44 111 47.44 117.92 87 109.58 87 109.58 67.38ZM107.08 87 107.08 106.43 99.25 106.43 99.25 76.07 96.75 76.07 96.75 106.38 88.1 106.38 88.2 69.88 107.08 69.88ZM117.94 106.43 109.58 106.43 109.58 89.45 117.94 89.45Z" fill="{color}"/></svg>''',
    
    # Realisering - Sykkel/fremdrift
    'realization': '''<svg viewBox="0 0 141.73 141.73" xmlns="http://www.w3.org/2000/svg"><path d="M116.62 80.1C116.61 89.49 108.99 97.09 99.6 97.08 90.21 97.07 82.61 89.45 82.62 80.06 82.63 74.37 85.48 69.06 90.23 65.91L99 83.59 101.23 82.47 92.42 64.69C94.67 63.63 97.12 63.08 99.6 63.08L99.6 60.58C96.73 60.58 93.89 61.22 91.3 62.46L85.3 50.46 85 49.88 88.77 49.88C91.53 49.88 93.77 52.12 93.77 54.88L93.77 56.13 96.27 56.13 96.27 54.86C96.26 50.72 92.91 47.37 88.77 47.36L80.87 47.36 83.09 51.57 88.67 62.79 62.27 62.79 58 47.87 45.84 47.87 45.84 50.37 56.09 50.37 59.9 63.59 55.29 67.37C47.26 60.19 34.92 60.88 27.74 68.92 20.56 76.95 21.25 89.29 29.29 96.47 37.32 103.65 49.66 102.95 56.84 94.92 59.75 91.66 61.48 87.52 61.76 83.16L68.14 83.16 63 65.29 86.92 65.29C78.74 72.3 77.78 84.61 84.79 92.8 91.8 100.98 104.12 101.94 112.3 94.93 116.63 91.22 119.12 85.8 119.12 80.1ZM42.3 98.93C32.91 98.91 25.32 91.27 25.35 81.88 25.37 72.5 33 64.9 42.39 64.93 46.38 64.94 50.25 66.36 53.3 68.93L41.67 78.52C40.57 79.47 40.44 81.14 41.39 82.24 41.89 82.82 42.62 83.16 43.38 83.16L59.26 83.16C58.62 92.05 51.21 98.93 42.3 98.93ZM64.81 80.66 43.38 80.66C43.38 80.66 43.29 80.66 43.25 80.57 43.21 80.48 43.25 80.45 43.25 80.44L60.65 66.21Z" fill="{color}"/></svg>''',
    
    # Realisert - Beholder/tank ferdig
    'realized': '''<svg viewBox="0 0 141.73 141.73" xmlns="http://www.w3.org/2000/svg"><path d="M59.87 37.12C59.87 32.98 63.23 29.62 67.37 29.62L68.62 29.62 68.62 32.12 67.37 32.12C64.61 32.12 62.37 34.36 62.37 37.12 62.37 39.88 64.61 42.12 67.37 42.12L68.62 42.12 68.62 44.62 67.37 44.62C63.23 44.62 59.87 41.26 59.87 37.12ZM109.33 80.31C109.32 84.45 105.97 87.79 101.83 87.8L97.1 87.8 97.1 85.3 101.83 85.3C104.59 85.3 106.83 83.06 106.83 80.3 106.83 77.54 104.59 75.3 101.83 75.3L93.83 75.3 93.83 77.3C93.83 90.55 83.08 101.3 69.83 101.3 56.58 101.3 45.83 90.55 45.83 77.3L45.83 63.4 93.83 63.4 93.83 72.81 101.83 72.81C105.97 72.82 109.32 76.17 109.33 80.31ZM91.33 65.9 48.33 65.9 48.33 77.32C48.33 89.19 57.96 98.82 69.83 98.82 81.7 98.82 91.33 89.19 91.33 77.32ZM45.84 107.61 93.84 107.61 93.84 105.11 45.84 105.11ZM79.84 49.61C79.83 45.47 76.48 42.12 72.34 42.11L71 42.11 71 44.61 72.25 44.61C75.01 44.61 77.25 46.85 77.25 49.61 77.25 52.37 75.01 54.61 72.25 54.61L71 54.61 71 57.11 72.25 57.11C76.4 57.13 79.77 53.78 79.79 49.63L79.79 49.61Z" fill="{color}"/></svg>''',
    
    # Personer - Interessenter/roller
    'people': '''<svg viewBox="0 0 141.73 141.73" xmlns="http://www.w3.org/2000/svg"><path d="M103.8 110.17 101.3 110.17 101.3 95.53C101.29 87.58 94.84 81.13 86.89 81.12L57 81.12C49.03 81.11 42.56 87.56 42.55 95.53L42.55 107.69 84.2 107.69 84.2 110.19 40.06 110.19 40.06 95.54C40.07 86.19 47.65 78.62 57 78.63L86.89 78.63C96.22 78.64 103.78 86.2 103.8 95.53ZM50.85 55.91 50.37 55.54 50.37 54.93C50.37 52.56 53.19 50.04 54.49 49.01 54.73 40.41 61.09 33.22 69.59 31.93 71.32 31.62 73.08 31.62 74.81 31.93 83.29 33.25 89.63 40.42 89.89 49 91.18 50 94 52.55 94 54.92L94 55.56 93.48 55.94C93.23 56.11 91.65 57.12 86.78 57.94L86.78 58.63C86.76 66.82 80.11 73.46 71.92 73.45 63.67 73.4 57.02 66.7 57.02 58.45L57.02 57.8C52.5 57.06 51.08 56.09 50.85 55.91ZM84.28 58.24C80.21 58.67 76.12 58.87 72.03 58.84 67.86 58.88 63.69 58.67 59.55 58.2L59.55 58.49C59.51 65.36 65.05 70.96 71.92 71L71.92 71C78.75 71 84.28 65.47 84.28 58.64L84.28 58.59ZM53.06 54.24C54.69 54.89 59.68 56.32 72 56.32 84.34 56.32 89.52 54.88 91.27 54.21 90.38 52.8 89.23 51.58 87.87 50.61L87.36 50.23 87.36 49.42C87.34 44.3 84.76 39.53 80.48 36.72L80.48 45.1 78 45.1 78 35.41C76.85 34.93 75.64 34.59 74.4 34.41L74.33 34.41C73.94 34.33 73.55 34.28 73.16 34.25L73.16 49.31 70.66 49.31 70.66 34.31C70.42 34.31 70.18 34.37 69.98 34.41 68.48 34.63 67.03 35.07 65.67 35.73L65.67 45.13 63.17 45.13 63.17 37.23C59.3 40.1 57.02 44.62 57 49.44L57 50.26 56.49 50.63C55.12 51.6 53.95 52.84 53.06 54.26Z" fill="{color}"/></svg>''',
    
    # Milepæler/prosess
    'milestones': '''<svg viewBox="0 0 141.73 141.73" xmlns="http://www.w3.org/2000/svg"><path d="M51.65 69.24C42.96 69.16 35.86 76.14 35.78 84.83 35.71 93.14 42.11 100.07 50.4 100.66L50.4 112.66 52.9 112.66 52.9 100.66C61.57 100.05 68.09 92.52 67.48 83.86 66.89 75.57 59.96 69.17 51.65 69.24ZM51.65 98.24C44.34 98.24 38.41 92.31 38.41 85 38.41 77.69 44.34 71.76 51.65 71.76 58.96 71.76 64.89 77.69 64.89 85 64.87 92.3 58.95 98.21 51.65 98.22ZM60.24 85.31 57.74 85.31C57.73 82.24 55.25 79.76 52.18 79.75L52.18 77.25C56.62 77.26 60.22 80.85 60.24 85.29ZM87.79 44.88C87.8 36.19 80.76 29.13 72.07 29.12 63.38 29.11 56.32 36.15 56.31 44.84 56.3 53.05 62.6 59.89 70.79 60.55L70.79 112.67 73.29 112.67 73.29 60.55C81.47 59.9 87.77 53.08 87.79 44.88ZM72.05 58.12C64.74 58.12 58.81 52.19 58.81 44.88 58.81 37.57 64.74 31.64 72.05 31.64 79.36 31.64 85.29 37.57 85.29 44.88 85.28 52.19 79.36 58.11 72.05 58.12ZM80.64 45.18 78.14 45.18C78.13 42.11 75.65 39.63 72.58 39.62L72.58 37.12C77.03 37.13 80.63 40.73 80.64 45.18ZM93.47 60.35C84.78 60.27 77.68 67.25 77.6 75.94 77.53 84.25 83.93 91.18 92.22 91.77L92.22 112.67 94.72 112.67 94.72 91.77C103.39 91.16 109.91 83.63 109.3 74.97 108.71 66.68 101.78 60.28 93.47 60.35ZM93.47 89.35C86.16 89.35 80.23 83.42 80.23 76.11 80.23 68.8 86.16 62.87 93.47 62.87 100.78 62.87 106.71 68.8 106.71 76.11 106.69 83.41 100.77 89.32 93.47 89.33ZM102.06 76.42 99.56 76.42C99.57 73.34 97.08 70.85 94 70.84L94 68.34C98.45 68.35 102.05 71.95 102.06 76.4Z" fill="{color}"/></svg>''',
    
    # Intervju/samarbeid
    'interview': '''<svg viewBox="0 0 141.73 141.73" xmlns="http://www.w3.org/2000/svg"><path d="M112.79 92 112.79 47.17 28 47.17 28 112.29 112.75 112.29 112.75 109.79 72.1 109.79 72.1 49.67 110.29 49.67 110.29 92ZM69.6 109.79 30.54 109.79 30.54 49.67 69.6 49.67ZM38.9 38.23 37.21 36.4 38.12 35.55C42 32 42 32 42.34 31.62 42.68 31.24 43.51 30.47 48.93 25.37L49.77 24.57 50.63 25.36C56.08 30.36 57.49 31.59 58.16 32.2 58.83 32.81 58.85 32.81 61.35 35.2L62.56 36.34 60.84 38.16 59.64 37C57.18 34.68 57.18 34.68 56.5 34.08 55.82 33.48 54.5 32.35 49.8 28 45.16 32.37 44.45 33.08 44.09 33.42 43.73 33.76 43.69 33.82 39.82 37.42ZM79.63 28.3 78.42 27.16 80.14 25.34 81.34 26.48C83.8 28.8 83.8 28.8 84.48 29.4 85.16 30 86.48 31.13 91.18 35.48 95.82 31.11 96.53 30.4 96.89 30.06 97.25 29.72 97.29 29.66 101.16 26.06L102.08 25.21 103.78 27.04 102.86 27.89C99 31.47 99 31.47 98.64 31.84 98.28 32.21 97.47 33 92.05 38.09L91.21 38.89 90.35 38.1C84.9 33.1 83.49 31.87 82.82 31.26 82.15 30.65 82.13 30.65 79.63 28.3ZM91.1 68.64C94.88 68.64 97.94 65.58 97.94 61.8 97.94 58.02 94.88 54.96 91.1 54.96 87.32 54.96 84.26 58.02 84.26 61.8 84.27 65.58 87.32 68.63 91.1 68.64ZM91.1 57.46C93.5 57.46 95.44 59.4 95.44 61.8 95.44 64.2 93.5 66.14 91.1 66.14 88.7 66.14 86.76 64.2 86.76 61.8 86.77 59.41 88.71 57.47 91.1 57.46ZM97.73 104.1 97.73 82C97.73 78.34 94.76 75.37 91.1 75.37 87.44 75.37 84.47 78.34 84.47 82L84.47 104.1 82 104.1 82 82C82 76.96 86.09 72.87 91.13 72.87 96.17 72.87 100.26 76.96 100.26 82L100.26 104.1ZM49.31 68.64C53.09 68.64 56.15 65.58 56.15 61.8 56.15 58.02 53.09 54.96 49.31 54.96 45.53 54.96 42.47 58.02 42.47 61.8 42.47 65.58 45.53 68.64 49.31 68.64ZM49.31 57.46C51.71 57.44 53.66 59.37 53.68 61.77 53.7 64.17 51.77 66.12 49.37 66.14 46.97 66.16 45.02 64.23 45 61.83L45 61.8C45 59.41 46.92 57.48 49.31 57.46ZM42.68 104.1 40.18 104.1 40.18 81.32C40.18 76.83 44.45 72.89 49.31 72.89 54.17 72.89 58.44 76.83 58.44 81.32L58.44 104.1 55.94 104.1 55.94 81.32C55.94 78.52 53.1 75.39 49.31 75.39 45.52 75.39 42.68 78.52 42.68 81.32Z" fill="{color}"/></svg>''',
    
    # Dashboard/organisasjon
    'dashboard': '''<svg viewBox="0 0 141.73 141.73" xmlns="http://www.w3.org/2000/svg"><path d="M28.58 109.76 66.65 109.76 66.65 73.38 28.58 73.38ZM31.08 75.88 64.15 75.88 64.15 107.26 31.08 107.26ZM40.18 95.46 55.05 95.46 55.05 88 40.18 88ZM42.68 90.46 52.55 90.46 52.55 93 42.68 93ZM74.54 73.38 74.54 109.76 112.61 109.76 112.61 73.38ZM110.11 107.26 77 107.26 77 75.88 110.07 75.88ZM100 88 85.14 88 85.14 95.47 100 95.47ZM97.5 93 87.64 93 87.64 90.49 97.5 90.49ZM90.5 29.54 52.43 29.54 52.43 65.88 90.5 65.88ZM88 63.38 54.93 63.38 54.93 32 88 32ZM78.9 44.11 64 44.11 64 51.58 78.9 51.58ZM76.4 49.11 66.53 49.11 66.53 46.61 76.4 46.61Z" fill="{color}"/></svg>''',
    
    # Mål/target
    'target': '''<svg viewBox="0 0 141.73 141.73" xmlns="http://www.w3.org/2000/svg"><path d="M121.18 91.28 118.72 91.28C118.71 85.94 115.15 81.27 110 79.85L73 69.77 71.74 69.77 36 78.33C30.45 79.65 26.37 84.37 25.86 90.05L101.62 90.05 101.62 92.51 23.34 92.51 23.34 91.28C23.33 83.98 28.34 77.63 35.44 75.94L71.12 67.39 71.12 64C65.64 63.34 61.72 58.36 62.38 52.88 63.04 47.4 68.02 43.48 73.5 44.14 78.54 44.75 82.33 49.03 82.31 54.11L79.85 54.11C79.85 49.97 76.49 46.61 72.35 46.61 68.21 46.61 64.85 49.97 64.85 54.11 64.85 58.25 68.21 61.61 72.35 61.61L73.58 61.61 73.58 67.41 110.63 77.51C116.84 79.21 121.16 84.84 121.18 91.28Z" fill="{color}"/></svg>'''
}

def get_icon_svg(icon_name, color=None, size=24):
    """Returner SVG-ikon med spesifisert farge og størrelse"""
    if color is None:
        color = COLORS['primary_dark']
    if icon_name in ICONS:
        svg = ICONS[icon_name].format(color=color)
        return f'<div style="width:{size}px;height:{size}px;display:inline-block;vertical-align:middle;">{svg}</div>'
    return ""

# Fase-ikoner mapping
PHASE_ICONS = {
    "Planlegging": "planning",
    "Gjennomføring": "execution",
    "Realisering": "realization",
    "Realisert": "realized"
}

# ============================================================================
# HENSIKT OG FORMÅL
# ============================================================================
HENSIKT_TEKST = """
### Hensikt
Modenhetsvurderingen har som formål å synliggjøre gode erfaringer og identifisere forbedringsområder i vårt arbeid med gevinster. Vi ønsker å lære av hverandre, dele beste praksis og hjelpe initiativer til å lykkes bedre med å skape og realisere gevinster.

Et sentralt fokusområde er å sikre at gevinstene vi arbeider med er konkrete og realitetsorienterte. Dette innebærer at nullpunkter og estimater er testet og validert, at hypoteser er prøvd mot representative caser og faktiske arbeidsforhold, og at forutsetningene for gevinstuttak er realistiske og forankret.

### Hvem inviteres?
Vi ønsker å intervjue alle som har vært eller er involvert i gevinstarbeidet – enten du har bidratt til utarbeidelse av business case, gevinstkart, gevinstrealiseringsplaner eller målinger, eller du har hatt ansvar for oppfølging og realisering.

### Hva vurderes?
Intervjuene dekker hele gevinstlivssyklusen – fra planlegging og gjennomføring til realisering og evaluering.

### Gevinster i endringsinitiativ
Et endringsinitiativ kan ha flere konkrete gevinster. Intervjuene kan gjennomføres med fokus på én spesifikk gevinst, eller for initiativet som helhet.
"""

# ============================================================================
# ROLLEDEFINISJONER
# ============================================================================
ROLES = {
    "Prosjektleder / Programleder": {
        "description": "Ansvar for overordnet gjennomføring og leveranser",
        "recommended_questions": {
            "Planlegging": [1, 2, 3, 4, 8, 9, 14, 16, 17, 18, 19, 20, 21, 22, 23],
            "Gjennomføring": [1, 2, 3, 4, 6, 8, 9, 14, 16, 17, 18, 19, 20, 21, 22, 23],
            "Realisering": [1, 2, 3, 8, 16, 17, 18, 19, 20, 22, 23],
            "Realisert": [1, 2, 3, 16, 17, 18, 19, 20, 22, 23]
        }
    },
    "Gevinsteier": {
        "description": "Ansvar for at gevinster realiseres i linjen",
        "recommended_questions": {
            "Planlegging": [2, 3, 4, 6, 9, 10, 11, 12, 13, 16, 17, 20, 21],
            "Gjennomføring": [2, 6, 9, 10, 11, 12, 13, 16, 17, 20, 21],
            "Realisering": [1, 2, 3, 6, 8, 9, 10, 11, 12, 13, 16, 17, 20, 21],
            "Realisert": [1, 2, 6, 8, 11, 12, 13, 16, 17, 20, 21]
        }
    },
    "Gevinstansvarlig": {
        "description": "Operativt ansvar for oppfølging og rapportering av gevinster",
        "recommended_questions": {
            "Planlegging": [1, 3, 5, 6, 7, 9, 10, 11, 12, 13, 14, 15, 16, 17],
            "Gjennomføring": [1, 5, 6, 7, 9, 10, 11, 12, 13, 14, 15, 16, 17],
            "Realisering": [1, 5, 6, 7, 9, 10, 11, 12, 13, 14, 15, 16, 17],
            "Realisert": [1, 5, 6, 7, 9, 10, 11, 12, 13, 14, 15, 16, 17]
        }
    },
    "Linjeleder / Mottaker": {
        "description": "Skal ta imot endringer og realisere gevinster i drift",
        "recommended_questions": {
            "Planlegging": [2, 8, 9, 12, 13, 18, 19, 20, 22],
            "Gjennomføring": [2, 8, 9, 12, 13, 17, 18, 19, 20, 22],
            "Realisering": [1, 2, 8, 9, 12, 13, 17, 18, 19, 20, 22],
            "Realisert": [1, 2, 8, 12, 13, 17, 18, 19, 20, 22]
        }
    },
    "Business Case-ansvarlig": {
        "description": "Utarbeidet gevinstgrunnlag og estimater",
        "recommended_questions": {
            "Planlegging": [1, 2, 3, 4, 5, 6, 7, 9, 10, 11, 12, 13, 14, 15, 21],
            "Gjennomføring": [1, 5, 6, 7, 10, 11, 14, 15, 21],
            "Realisering": [1, 5, 6, 7, 10, 11, 14, 15],
            "Realisert": [1, 5, 6, 7, 10, 11, 14, 15, 21]
        }
    },
    "Styringsgruppe": {
        "description": "Overordnet ansvar og beslutninger",
        "recommended_questions": {
            "Planlegging": [2, 4, 8, 14, 16, 19, 20, 21, 22],
            "Gjennomføring": [2, 4, 8, 14, 16, 19, 20, 22],
            "Realisering": [2, 4, 8, 16, 19, 20, 22],
            "Realisert": [2, 4, 8, 16, 19, 20, 22]
        }
    },
    "Controller / Økonomi": {
        "description": "Oppfølging av økonomiske gevinster",
        "recommended_questions": {
            "Planlegging": [2, 4, 5, 6, 11, 12, 13, 21],
            "Gjennomføring": [2, 5, 6, 11, 12, 13, 21],
            "Realisering": [2, 5, 6, 11, 12, 13, 21],
            "Realisert": [2, 5, 6, 11, 12, 13, 21]
        }
    },
    "Endringsleder": {
        "description": "Ansvar for endringsledelse og kommunikasjon",
        "recommended_questions": {
            "Planlegging": [2, 8, 9, 18, 19, 20, 22, 23],
            "Gjennomføring": [2, 8, 9, 18, 19, 20, 22, 23],
            "Realisering": [1, 2, 8, 9, 18, 19, 20, 22, 23],
            "Realisert": [1, 2, 8, 18, 19, 20, 22, 23]
        }
    }
}

# ============================================================================
# PARAMETERE
# ============================================================================
PARAMETERS = {
    "Strategisk forankring": {
        "description": "Strategisk retning, kobling til mål og KPI-er",
        "questions": [2, 4]
    },
    "Gevinstkart og visualisering": {
        "description": "Gevinstkart, sammenhenger mellom tiltak og effekter",
        "questions": [3]
    },
    "Nullpunkter og estimater": {
        "description": "Kvalitet på nullpunkter, estimater og datagrunnlag",
        "questions": [6, 7, 11]
    },
    "Interessenter og forankring": {
        "description": "Interessentengasjement, kommunikasjon og forankring",
        "questions": [8, 19]
    },
    "Eierskap og ansvar": {
        "description": "Roller, ansvar og eierskap for gevinstuttak",
        "questions": [20]
    },
    "Forutsetninger og risiko": {
        "description": "Gevinstforutsetninger, risiko og ulemper",
        "questions": [9, 10, 14, 15]
    },
    "Gevinstrealiseringsplan": {
        "description": "Plan som operativt styringsverktøy",
        "questions": [16, 17]
    },
    "Effektivitet og produktivitet": {
        "description": "Måling, disponering og bærekraft",
        "questions": [12, 13]
    },
    "Læring og forbedring": {
        "description": "Bruk av tidligere erfaringer og kontinuerlig læring",
        "questions": [1]
    },
    "Momentum og tidlig gevinstuttak": {
        "description": "Bygge momentum gjennom tidlig gevinstrealisering",
        "questions": [5, 21, 22, 23]
    }
}

# ============================================================================
# KOMPLETT SPØRSMÅLSSETT - ALLE 23 SPØRSMÅL PER FASE (ORIGINALE FULLSTENDIGE)
# ============================================================================
PHASES = ["Planlegging", "Gjennomføring", "Realisering", "Realisert"]

questions_data = {
    "Planlegging": [
        {
            "id": 1,
            "title": "Bruk av tidligere læring og gevinstdata",
            "question": "Hvordan anvendes erfaringer og læring fra tidligere prosjekter og gevinstarbeid i planleggingen av nye gevinster?",
            "scale": [
                "Nivå 1: Ingen læring fra tidligere arbeid anvendt.",
                "Nivå 2: Enkelte erfaringer omtalt, men ikke strukturert brukt.",
                "Nivå 3: Læring inkludert i planlegging for enkelte områder.",
                "Nivå 4: Systematisk bruk av tidligere gevinstdata i planlegging og estimering.",
                "Nivå 5: Kontinuerlig læring integrert i planleggingsprosessen og gevinststrategien."
            ]
        },
        {
            "id": 2,
            "title": "Strategisk retning og gevinstforståelse",
            "question": "Hvilke gevinster arbeider dere med, og hvorfor er de viktige for organisasjonens strategiske mål?",
            "scale": [
                "Nivå 1: Gevinster er vagt definert, uten tydelig kobling til strategi.",
                "Nivå 2: Gevinster er identifisert, men mangler klare kriterier og prioritering.",
                "Nivå 3: Gevinster er dokumentert og delvis knyttet til strategiske mål, men grunnlaget har usikkerhet.",
                "Nivå 4: Gevinster er tydelig koblet til strategiske mål med konkrete måltall.",
                "Nivå 5: Gevinster er fullt integrert i styringssystemet og brukes i beslutninger."
            ]
        },
        {
            "id": 3,
            "title": "Gevinstkart og visualisering",
            "question": "Er gevinstene synliggjort i gevinstkartet, med tydelig sammenheng mellom tiltak, effekter og mål?",
            "scale": [
                "Nivå 1: Gevinstkart finnes ikke eller er utdatert.",
                "Nivå 2: Et foreløpig gevinstkart eksisterer, men dekker ikke hele området.",
                "Nivå 3: Kartet inkluderer hovedgevinster, men mangler validering og detaljer.",
                "Nivå 4: Kartet er brukt aktivt i planlegging og oppfølging.",
                "Nivå 5: Gevinstkartet oppdateres kontinuerlig og er integrert i styringsdialoger."
            ]
        },
        {
            "id": 4,
            "title": "Strategisk kobling og KPI-er",
            "question": "Er gevinstene tydelig knyttet til strategiske mål og eksisterende KPI-er?",
            "scale": [
                "Nivå 1: Ingen kobling mellom gevinster og strategi eller KPI-er.",
                "Nivå 2: Kobling er antatt, men ikke dokumentert.",
                "Nivå 3: Kobling er etablert for enkelte KPI-er, men ikke konsistent.",
                "Nivå 4: Tydelig kobling mellom gevinster og relevante KPI-er.",
                "Nivå 5: Koblingen følges opp i styringssystem og rapportering."
            ]
        },
        {
            "id": 5,
            "title": "Avgrensning av programgevinst",
            "question": "Er det tydelig avklart hvilke effekter som stammer fra programmet versus andre tiltak eller økte rammer?",
            "scale": [
                "Nivå 1: Ingen skille mellom program- og eksterne effekter.",
                "Nivå 2: Delvis omtalt, men uklart hva som er innenfor programmet.",
                "Nivå 3: Avgrensning er gjort i plan, men ikke dokumentert grundig.",
                "Nivå 4: Avgrensning er dokumentert og anvendt i beregninger.",
                "Nivå 5: Effektisolering er standard praksis og brukes systematisk."
            ]
        },
        {
            "id": 6,
            "title": "Nullpunkter og estimater",
            "question": "Er nullpunkter og estimater etablert, testet og dokumentert på en konsistent og troverdig måte?",
            "scale": [
                "Nivå 1: Nullpunkter mangler eller bygger på uprøvde antagelser.",
                "Nivå 2: Enkelte nullpunkter finnes, men uten felles metode.",
                "Nivå 3: Nullpunkter og estimater er definert, men med høy usikkerhet.",
                "Nivå 4: Nullpunkter og estimater er basert på testede data og validerte metoder.",
                "Nivå 5: Nullpunkter og estimater kvalitetssikres jevnlig og brukes aktivt til læring."
            ]
        },
        {
            "id": 7,
            "title": "Hypotesetesting og datagrunnlag",
            "question": "Finnes formell prosess for hypotesetesting på representative caser?",
            "scale": [
                "Nivå 1: Ikke etablert/uklart; ingen dokumenterte praksiser.",
                "Nivå 2: Delvis definert; uformell praksis uten forankring/validering.",
                "Nivå 3: Etablert for deler av området; variabel kvalitet.",
                "Nivå 4: Godt forankret og systematisk anvendt; måles og følges opp.",
                "Nivå 5: Fullt integrert i styring; kontinuerlig forbedring og læring."
            ]
        },
        {
            "id": 8,
            "title": "Interessentengasjement",
            "question": "Ble relevante interessenter involvert i utarbeidelsen av gevinstgrunnlag?",
            "scale": [
                "Nivå 1: Ingen involvering av interessenter.",
                "Nivå 2: Begrenset og ustrukturert involvering.",
                "Nivå 3: Bred deltakelse, men uten systematisk prosess.",
                "Nivå 4: Systematisk og koordinert involvering med klar rollefordeling.",
                "Nivå 5: Kontinuerlig engasjement med dokumentert medvirkning."
            ]
        },
        {
            "id": 9,
            "title": "Gevinstforutsetninger",
            "question": "Er alle vesentlige forutsetninger ivaretatt for å muliggjøre gevinstrealisering?",
            "scale": [
                "Nivå 1: Ingen kartlegging av gevinstforutsetninger.",
                "Nivå 2: Noen forutsetninger er identifisert, men ikke systematisk dokumentert.",
                "Nivå 3: Hovedforutsetninger er dokumentert, men uten klar eierskap.",
                "Nivå 4: Alle kritiske forutsetninger er kartlagt med tildelt ansvar.",
                "Nivå 5: Gevinstforutsetninger er integrert i risikostyring og oppfølges kontinuerlig."
            ]
        },
        {
            "id": 10,
            "title": "Prinsipielle og vilkårsmessige kriterier",
            "question": "Er forutsetninger og kriterier som påvirker gevinstene tydelig definert og dokumentert?",
            "scale": [
                "Nivå 1: Ingen kriterier dokumentert.",
                "Nivå 2: Kriterier er beskrevet uformelt.",
                "Nivå 3: Kriterier dokumentert i deler av planverket.",
                "Nivå 4: Vesentlige kriterier er analysert og håndtert i gevinstrealiseringsplanen.",
                "Nivå 5: Kriterier overvåkes, følges opp og inngår i risikostyringen."
            ]
        },
        {
            "id": 11,
            "title": "Enighet om nullpunkter/estimater",
            "question": "Er det oppnådd enighet blant nøkkelinteressenter om nullpunkter og estimater?",
            "scale": [
                "Nivå 1: Ingen enighet eller dokumentert praksis.",
                "Nivå 2: Delvis enighet, men ikke formalisert.",
                "Nivå 3: Enighet for hovedestimater, men med reservasjoner.",
                "Nivå 4: Full enighet dokumentert og forankret.",
                "Nivå 5: Kontinuerlig dialog og justering av estimater med interessentene."
            ]
        },
        {
            "id": 12,
            "title": "Disponering av kostnads- og tidsbesparelser",
            "question": "Hvordan er kostnads- og tidsbesparelser planlagt disponert mellom prissatte og ikke-prissatte gevinster?",
            "scale": [
                "Nivå 1: Ingen plan for disponering eller måling av besparelser.",
                "Nivå 2: Delvis oversikt, men ikke dokumentert eller fulgt opp.",
                "Nivå 3: Plan finnes for enkelte områder, men uten systematikk.",
                "Nivå 4: Disponering og effekter dokumentert og målt.",
                "Nivå 5: Frigjorte ressurser disponeres strategisk og måles som del av gevinstrealiseringen."
            ]
        },
        {
            "id": 13,
            "title": "Måling av effektivitet og produktivitet",
            "question": "Hvordan måles økt effektivitet og produktivitet som følge av besparelser?",
            "scale": [
                "Nivå 1: Ingen måling av effektivitet eller produktivitet.",
                "Nivå 2: Enkelte målinger, men ikke systematisk.",
                "Nivå 3: Måling for enkelte gevinster, men begrenset fokus på bærekraft.",
                "Nivå 4: Systematisk måling og vurdering av om gevinster opprettholdes over tid.",
                "Nivå 5: Måling integrert i gevinstoppfølgingen, bærekraftige gevinster sikres."
            ]
        },
        {
            "id": 14,
            "title": "Operasjonell risiko og ulemper",
            "question": "Er mulige negative konsekvenser eller ulemper knyttet til operasjonelle forhold identifisert, vurdert og håndtert i planen?",
            "scale": [
                "Nivå 1: Negative effekter ikke vurdert.",
                "Nivå 2: Kjent, men ikke håndtert.",
                "Nivå 3: Beskrevet, men ikke fulgt opp systematisk.",
                "Nivå 4: Håndtert og overvåket med tilpasning til ulike operasjonelle scenarier.",
                "Nivå 5: Systematisk vurdert og del av gevinstdialogen med kontinuerlig justering."
            ]
        },
        {
            "id": 15,
            "title": "Balanse mellom gevinster og ulemper",
            "question": "Hvordan sikres det at balansen mellom gevinster og ulemper vurderes i styringsdialoger?",
            "scale": [
                "Nivå 1: Ingen vurdering av balanse.",
                "Nivå 2: Diskuteres uformelt.",
                "Nivå 3: Del av enkelte oppfølgingsmøter.",
                "Nivå 4: Systematisk vurdert i gevinststyring.",
                "Nivå 5: Inngår som fast punkt i styrings- og gevinstdialoger."
            ]
        },
        {
            "id": 16,
            "title": "Dokumentasjon og gevinstrealiseringsplan",
            "question": "Er det utarbeidet en forankret gevinstrealiseringsplan som beskriver hvordan gevinstene skal hentes ut og måles?",
            "scale": [
                "Nivå 1: Ingen formell gevinstrealiseringsplan.",
                "Nivå 2: Utkast til plan finnes, men er ufullstendig.",
                "Nivå 3: Plan er etablert, men ikke validert eller periodisert.",
                "Nivå 4: Planen er forankret, oppdatert og koblet til gevinstkartet.",
                "Nivå 5: Planen brukes aktivt som styringsdokument med revisjon."
            ]
        },
        {
            "id": 17,
            "title": "Gevinstrealiseringsplan som operativ handlingsplan",
            "question": "Hvordan sikres det at gevinstrealiseringsplanen fungerer som en operativ handlingsplan i linjen med tilpasning til ulike strekningsforhold?",
            "scale": [
                "Nivå 1: Planen brukes ikke som operativt styringsverktøy.",
                "Nivå 2: Plan finnes, men uten operativ oppfølging.",
                "Nivå 3: Planen følges delvis opp i linjen.",
                "Nivå 4: Planen brukes aktivt som handlingsplan og styringsverktøy.",
                "Nivå 5: Gevinstplanen er fullt operativt integrert i linjens handlingsplaner og rapportering med tilpasning til lokale forhold."
            ]
        },
        {
            "id": 18,
            "title": "Endringsberedskap og operativ mottaksevne",
            "question": "Er organisasjonen forberedt og har den tilstrekkelig kapasitet til å ta imot endringer og nye arbeidsformer som følger av programmet, inkludert tilpasning til ulike strekningsforhold?",
            "scale": [
                "Nivå 1: Ingen plan for endringsberedskap.",
                "Nivå 2: Kapasitet vurderes uformelt, men ikke håndtert.",
                "Nivå 3: Endringskapasitet omtales, men uten konkrete tiltak.",
                "Nivå 4: Tilfredsstillende beredskap etablert og koordinert med linjen.",
                "Nivå 5: Endringskapasitet er strukturert, overvåket og integrert i styring med tilpasning til lokale forhold."
            ]
        },
        {
            "id": 19,
            "title": "Kommunikasjon og forankring",
            "question": "Er gevinstgrunnlag, roller og forventninger godt kommunisert i organisasjonen?",
            "scale": [
                "Nivå 1: Ingen felles forståelse eller kommunikasjon.",
                "Nivå 2: Informasjon deles sporadisk.",
                "Nivå 3: Kommunikasjon er planlagt, men ikke systematisk målt.",
                "Nivå 4: Kommunikasjon er systematisk og forankret i organisasjonen.",
                "Nivå 5: Forankring skjer løpende som del av styringsdialog."
            ]
        },
        {
            "id": 20,
            "title": "Eierskap og ansvar",
            "question": "Er ansvar og roller tydelig definert for å sikre gjennomføring og gevinstuttak?",
            "scale": [
                "Nivå 1: Ansvar er uklart eller mangler.",
                "Nivå 2: Ansvar er delvis definert, men ikke praktisert.",
                "Nivå 3: Ansvar er kjent, men samhandling varierer.",
                "Nivå 4: Roller og ansvar fungerer godt i praksis.",
                "Nivå 5: Sterkt eierskap og kultur for ansvarliggjøring."
            ]
        },
        {
            "id": 21,
            "title": "Periodisering og forankring",
            "question": "Er gevinstrealiseringsplanen periodisert, validert og godkjent av ansvarlige?",
            "scale": [
                "Nivå 1: Ingen tidsplan eller forankring.",
                "Nivå 2: Tidsplan foreligger, men ikke validert.",
                "Nivå 3: Delvis forankret hos enkelte ansvarlige/eiere.",
                "Nivå 4: Fullt forankret og koordinert med budsjett- og styringsprosesser.",
                "Nivå 5: Planen brukes aktivt i styringsdialog og rapportering."
            ]
        },
        {
            "id": 22,
            "title": "Realisme og engasjement",
            "question": "Opplever dere at gevinstplanen og estimatene oppleves realistiske og engasjerer eierne og interessentene?",
            "scale": [
                "Nivå 1: Ingen troverdighet eller engasjement.",
                "Nivå 2: Begrenset tillit til estimater.",
                "Nivå 3: Delvis aksept, men varierende engasjement.",
                "Nivå 4: Høy troverdighet og engasjement.",
                "Nivå 5: Sterk troverdighet og aktiv motivasjon i organisasjonen."
            ]
        },
        {
            "id": 23,
            "title": "Bygge momentum og tidlig gevinstuttak",
            "question": "Hvordan planlegges det for å bygge momentum og realisere tidlige gevinster underveis i programmet?",
            "scale": [
                "Nivå 1: Ingen plan for tidlig gevinstuttak eller oppbygging av momentum.",
                "Nivå 2: Enkelte uformelle vurderinger av tidlige gevinster.",
                "Nivå 3: Plan for tidlig gevinstuttak er identifisert, men ikke koordinert.",
                "Nivå 4: Strukturert tilnærming for tidlig gevinstuttak med tildelt ansvar.",
                "Nivå 5: Tidlig gevinstuttak er integrert i programmets DNA og brukes aktivt for å bygge momentum."
            ]
        }
    ],
    "Gjennomføring": [
        {
            "id": 1,
            "title": "Bruk av tidligere læring og gevinstdata",
            "question": "Hvordan brukes erfaringer og læring fra tidligere prosjekter og gevinstarbeid til å justere tiltak under gjennomføringen?",
            "scale": [
                "Nivå 1: Ingen læring fra tidligere arbeid anvendt under gjennomføring.",
                "Nivå 2: Enkelte erfaringer omtalt, men ikke strukturert brukt for justering.",
                "Nivå 3: Læring inkludert i justering for enkelte områder under gjennomføring.",
                "Nivå 4: Systematisk bruk av tidligere gevinstdata for å justere tiltak underveis.",
                "Nivå 5: Kontinuerlig læring integrert i gjennomføringsprosessen og gevinstjustering."
            ]
        },
        {
            "id": 2,
            "title": "Strategisk retning og gevinstforståelse",
            "question": "Hvordan opprettholdes den strategiske retningen og forståelsen av gevinster under gjennomføring?",
            "scale": [
                "Nivå 1: Strategisk kobling glemmes under gjennomføring.",
                "Nivå 2: Strategi omtales, men ikke operasjonalisert i gjennomføring.",
                "Nivå 3: Strategisk kobling vedlikeholdes i deler av gjennomføringen.",
                "Nivå 4: Tydelig strategisk retning i gjennomføring med regelmessig oppdatering.",
                "Nivå 5: Strategi og gevinstforståelse dynamisk tilpasses underveis basert på læring."
            ]
        },
        {
            "id": 3,
            "title": "Gevinstkart og visualisering",
            "question": "Hvordan brukes gevinstkartet aktivt under gjennomføring for å styre og kommunisere fremdrift?",
            "scale": [
                "Nivå 1: Gevinstkartet brukes ikke under gjennomføring.",
                "Nivå 2: Gevinstkartet vises, men ikke aktivt brukt.",
                "Nivå 3: Gevinstkartet oppdateres og brukes i noen beslutninger.",
                "Nivå 4: Gevinstkartet er aktivt styringsverktøy under gjennomføring.",
                "Nivå 5: Gevinstkartet brukes dynamisk til å justere strategi og tiltak underveis."
            ]
        },
        {
            "id": 4,
            "title": "Strategisk kobling og KPI-er",
            "question": "Hvordan følges opp den strategiske koblingen og KPI-ene under gjennomføring?",
            "scale": [
                "Nivå 1: Ingen oppfølging av strategisk kobling under gjennomføring.",
                "Nivå 2: KPI-er måles, men kobling til strategi mangler.",
                "Nivå 3: Noen KPI-er følges opp med strategisk kobling.",
                "Nivå 4: Systematisk oppfølging av KPI-er med tydelig strategisk kobling.",
                "Nivå 5: Dynamisk justering av KPI-er basert på strategisk utvikling underveis."
            ]
        },
        {
            "id": 5,
            "title": "Avgrensning av programgevinst",
            "question": "Hvordan håndteres avgrensning av programgevinster under gjennomføring når nye forhold oppstår?",
            "scale": [
                "Nivå 1: Avgrensning glemmes under gjennomføring.",
                "Nivå 2: Avgrensning omtales, men ikke operasjonalisert.",
                "Nivå 3: Avgrensning håndteres for større endringer.",
                "Nivå 4: System for å håndtere avgrensning under gjennomføring.",
                "Nivå 5: Dynamisk avgrensningshåndtering integrert i beslutningsprosesser."
            ]
        },
        {
            "id": 6,
            "title": "Nullpunkter og estimater",
            "question": "Hvordan justeres nullpunkter og estimater under gjennomføring basert på nye data og erfaringer?",
            "scale": [
                "Nivå 1: Nullpunkter og estimater justeres ikke under gjennomføring.",
                "Nivå 2: Justering skjer ad hoc uten struktur.",
                "Nivå 3: Systematisk justering for store avvik.",
                "Nivå 4: Regelmessig revisjon og justering av nullpunkter og estimater.",
                "Nivå 5: Kontinuerlig justering basert på realtidsdata og læring."
            ]
        },
        {
            "id": 7,
            "title": "Hypotesetesting og datagrunnlag",
            "question": "Hvordan testes hypoteser og datagrunnlag under gjennomføring for å validere tilnærmingen?",
            "scale": [
                "Nivå 1: Hypoteser testes ikke under gjennomføring.",
                "Nivå 2: Noen uformelle tester gjennomføres.",
                "Nivå 3: Formell testing for kritiske hypoteser.",
                "Nivå 4: Systematisk testing og validering under gjennomføring.",
                "Nivå 5: Kontinuerlig hypotesetesting integrert i læringsprosesser."
            ]
        },
        {
            "id": 8,
            "title": "Interessentengasjement",
            "question": "Hvordan opprettholdes interessentengasjement under gjennomføring?",
            "scale": [
                "Nivå 1: Interessentengasjement avtar under gjennomføring.",
                "Nivå 2: Begrenset engasjement for viktige beslutninger.",
                "Nivå 3: Regelmessig engasjement for større endringer.",
                "Nivå 4: Systematisk interessentoppfølging under gjennomføring.",
                "Nivå 5: Kontinuerlig dialog og samskaping med interessenter."
            ]
        },
        {
            "id": 9,
            "title": "Gevinstforutsetninger",
            "question": "Hvordan overvåkes og håndteres gevinstforutsetninger under gjennomføring?",
            "scale": [
                "Nivå 1: Forutsetninger overvåkes ikke under gjennomføring.",
                "Nivå 2: Noen forutsetninger overvåkes uformelt.",
                "Nivå 3: Systematisk overvåkning av kritiske forutsetninger.",
                "Nivå 4: Aktiv håndtering av endrede forutsetninger.",
                "Nivå 5: Forutsetningsstyring integrert i risikostyring og beslutninger."
            ]
        },
        {
            "id": 10,
            "title": "Prinsipielle og vilkårsmessige kriterier",
            "question": "Hvordan håndteres endringer i prinsipielle og vilkårsmessige kriterier under gjennomføring?",
            "scale": [
                "Nivå 1: Endringer i kriterier håndteres ikke.",
                "Nivå 2: Store endringer håndteres reaktivt.",
                "Nivå 3: System for å håndtere endringer i kriterier.",
                "Nivå 4: Proaktiv håndtering av endrede kriterier.",
                "Nivå 5: Dynamisk tilpasning til endrede kriterier i sanntid."
            ]
        },
        {
            "id": 11,
            "title": "Enighet om nullpunkter/estimater",
            "question": "Hvordan opprettholdes enighet om nullpunkter og estimater under gjennomføring?",
            "scale": [
                "Nivå 1: Enighet testes ikke under gjennomføring.",
                "Nivå 2: Enighet bekreftes ved store endringer.",
                "Nivå 3: Regelmessig bekreftelse av enighet.",
                "Nivå 4: Systematisk arbeid for å opprettholde enighet.",
                "Nivå 5: Kontinuerlig dialog og justering for å opprettholde enighet."
            ]
        },
        {
            "id": 12,
            "title": "Disponering av kostnads- og tidsbesparelser",
            "question": "Hvordan håndteres disponering av besparelser under gjennomføring?",
            "scale": [
                "Nivå 1: Disponering håndteres ikke under gjennomføring.",
                "Nivå 2: Disponering justeres for store avvik.",
                "Nivå 3: Systematisk revisjon av disponeringsplaner.",
                "Nivå 4: Dynamisk tilpasning av disponering basert på resultater.",
                "Nivå 5: Optimal disponering integrert i beslutningsstøtte."
            ]
        },
        {
            "id": 13,
            "title": "Måling av effektivitet og produktivitet",
            "question": "Hvordan måles og følges opp effektivitet og produktivitet under gjennomføring?",
            "scale": [
                "Nivå 1: Effektivitet og produktivitet måles ikke underveis.",
                "Nivå 2: Noen målinger registreres, men ikke analysert.",
                "Nivå 3: Systematisk måling med begrenset analyse.",
                "Nivå 4: Regelmessig analyse og justering basert på målinger.",
                "Nivå 5: Realtids overvåkning og proaktiv justering."
            ]
        },
        {
            "id": 14,
            "title": "Operasjonell risiko og ulemper",
            "question": "Hvordan identifiseres og håndteres nye operasjonelle risikoer og ulemper under gjennomføring?",
            "scale": [
                "Nivå 1: Nye risikoer identifiseres ikke underveis.",
                "Nivå 2: Store risikoer håndteres reaktivt.",
                "Nivå 3: Systematisk identifisering av nye risikoer.",
                "Nivå 4: Proaktiv håndtering av nye risikoer.",
                "Nivå 5: Risikostyring integrert i daglig drift."
            ]
        },
        {
            "id": 15,
            "title": "Balanse mellom gevinster og ulemper",
            "question": "Hvordan vurderes balansen mellom gevinster og ulemper under gjennomføring?",
            "scale": [
                "Nivå 1: Balansen vurderes ikke under gjennomføring.",
                "Nivå 2: Balansen vurderes ved store endringer.",
                "Nivå 3: Regelmessig vurdering av balansen.",
                "Nivå 4: Systematisk overvåkning av balansen.",
                "Nivå 5: Balansevurdering integrert i beslutningsprosesser."
            ]
        },
        {
            "id": 16,
            "title": "Dokumentasjon og gevinstrealiseringsplan",
            "question": "Hvordan oppdateres og brukes gevinstrealiseringsplanen under gjennomføring?",
            "scale": [
                "Nivå 1: Gevinstrealiseringsplanen oppdateres ikke.",
                "Nivå 2: Planen oppdateres ved store endringer.",
                "Nivå 3: Regelmessig oppdatering av planen.",
                "Nivå 4: Planen brukes aktivt i styring og beslutninger.",
                "Nivå 5: Dynamisk oppdatering og bruk av planen i sanntid."
            ]
        },
        {
            "id": 17,
            "title": "Gevinstrealiseringsplan som operativ handlingsplan",
            "question": "Hvordan fungerer gevinstrealiseringsplanen som operativ handlingsplan under gjennomføring?",
            "scale": [
                "Nivå 1: Planen brukes ikke som operativ handlingsplan.",
                "Nivå 2: Planen brukes til visse operasjoner.",
                "Nivå 3: Planen er integrert i deler av den operative styringen.",
                "Nivå 4: Planen er aktivt operativt styringsverktøy.",
                "Nivå 5: Planen er fullt integrert i alle operative beslutninger."
            ]
        },
        {
            "id": 18,
            "title": "Endringsberedskap og operativ mottaksevne",
            "question": "Hvordan utvikles endringsberedskap og operativ mottaksevne under gjennomføring?",
            "scale": [
                "Nivå 1: Endringsberedskap utvikles ikke underveis.",
                "Nivå 2: Begrenset fokus på endringsberedskap.",
                "Nivå 3: Systematisk arbeid med endringsberedskap.",
                "Nivå 4: Målrettet utvikling av mottaksevne.",
                "Nivå 5: Kontinuerlig tilpasning og læring i endringsprosessen."
            ]
        },
        {
            "id": 19,
            "title": "Kommunikasjon og forankring",
            "question": "Hvordan opprettholdes kommunikasjon og forankring under gjennomføring?",
            "scale": [
                "Nivå 1: Kommunikasjon avtar under gjennomføring.",
                "Nivå 2: Begrenset kommunikasjon om viktige endringer.",
                "Nivå 3: Regelmessig kommunikasjon om fremdrift.",
                "Nivå 4: Systematisk kommunikasjonsplan under gjennomføring.",
                "Nivå 5: Kontinuerlig dialog og tilbakemelding integrert i prosessen."
            ]
        },
        {
            "id": 20,
            "title": "Eierskap og ansvar",
            "question": "Hvordan utøves eierskap og ansvar under gjennomføring?",
            "scale": [
                "Nivå 1: Eierskap og ansvar svekkes under gjennomføring.",
                "Nivå 2: Begrenset eierskap i kritiske faser.",
                "Nivå 3: Tydelig eierskap for sentrale ansvarsområder.",
                "Nivå 4: Aktivt utøvd eierskap gjennom hele prosessen.",
                "Nivå 5: Sterk eierskapskultur som driver gjennomføring."
            ]
        },
        {
            "id": 21,
            "title": "Periodisering og forankring",
            "question": "Hvordan justeres periodisering og forankring under gjennomføring?",
            "scale": [
                "Nivå 1: Periodisering justeres ikke under gjennomføring.",
                "Nivå 2: Store justeringer i periodisering.",
                "Nivå 3: Regelmessig revisjon av periodisering.",
                "Nivå 4: Dynamisk tilpasning av periodisering.",
                "Nivå 5: Fleksibel periodisering integrert i styringssystemet."
            ]
        },
        {
            "id": 22,
            "title": "Realisme og engasjement",
            "question": "Hvordan opprettholdes realisme og engasjement under gjennomføring?",
            "scale": [
                "Nivå 1: Realisme og engasjement avtar.",
                "Nivå 2: Begrenset fokus på å opprettholde engasjement.",
                "Nivå 3: Arbeid med å opprettholde realisme og engasjement.",
                "Nivå 4: Systematisk arbeid for å styrke troverdighet.",
                "Nivå 5: Høy troverdighet og engasjement gjennom hele prosessen."
            ]
        },
        {
            "id": 23,
            "title": "Bygge momentum og tidlig gevinstuttak",
            "question": "Hvordan bygges momentum gjennom tidlig gevinstuttak under gjennomføringsfasen?",
            "scale": [
                "Nivå 1: Ingen fokus på momentum eller tidlig gevinstuttak.",
                "Nivå 2: Noen tidlige gevinster realiseres, men uten strategi.",
                "Nivå 3: Planlagt for tidlig gevinstuttak, men begrenset gjennomføring.",
                "Nivå 4: Systematisk arbeid med tidlig gevinstuttak for å bygge momentum.",
                "Nivå 5: Kontinuerlig fokus på momentum gjennom suksessiv gevinstrealisering."
            ]
        }
    ],
    "Realisering": [
        {
            "id": 1,
            "title": "Bruk av tidligere læring og gevinstdata",
            "question": "Hvordan anvendes læring fra tidligere prosjekter og gevinstarbeid for å optimalisere gevinstuttak under realiseringen?",
            "scale": [
                "Nivå 1: Ingen læring anvendt i realiseringsfasen.",
                "Nivå 2: Enkelte erfaringer tas i betraktning.",
                "Nivå 3: Systematisk bruk av læring for å optimalisere uttak.",
                "Nivå 4: Læring integrert i realiseringsprosessen.",
                "Nivå 5: Kontinuerlig læring og optimalisering under realisering."
            ]
        },
        {
            "id": 2,
            "title": "Strategisk retning og gevinstforståelse",
            "question": "Hvordan sikres strategisk retning og gevinstforståelse under realiseringen?",
            "scale": [
                "Nivå 1: Strategisk retning glemmes under realisering.",
                "Nivå 2: Strategi refereres til, men ikke operasjonalisert.",
                "Nivå 3: Tydelig strategisk retning i realiseringsarbeid.",
                "Nivå 4: Strategi dynamisk tilpasses under realisering.",
                "Nivå 5: Strategi og realisering fullt integrert og sammenvevd."
            ]
        },
        {
            "id": 3,
            "title": "Gevinstkart og visualisering",
            "question": "Hvordan brukes gevinstkartet for å styre realiseringsarbeidet?",
            "scale": [
                "Nivå 1: Gevinstkartet brukes ikke under realisering.",
                "Nivå 2: Gevinstkartet vises, men ikke aktivt brukt.",
                "Nivå 3: Gevinstkartet brukes til å prioritere realisering.",
                "Nivå 4: Gevinstkartet er aktivt styringsverktøy.",
                "Nivå 5: Gevinstkartet dynamisk oppdateres basert på realisering."
            ]
        },
        {
            "id": 4,
            "title": "Strategisk kobling og KPI-er",
            "question": "Hvordan følges opp strategisk kobling og KPI-er under realiseringen?",
            "scale": [
                "Nivå 1: Ingen oppfølging av strategisk kobling.",
                "Nivå 2: KPI-er måles, men kobling til strategi svak.",
                "Nivå 3: Systematisk oppfølging av strategisk kobling.",
                "Nivå 4: Dynamisk justering basert på KPI-utvikling.",
                "Nivå 5: Full integrasjon mellom strategi, KPI-er og realisering."
            ]
        },
        {
            "id": 5,
            "title": "Avgrensning av programgevinst",
            "question": "Hvordan håndteres avgrensning av programgevinster under realiseringen?",
            "scale": [
                "Nivå 1: Avgrensning håndteres ikke under realisering.",
                "Nivå 2: Store avgrensningsutfordringer håndteres.",
                "Nivå 3: System for å håndtere avgrensning.",
                "Nivå 4: Proaktiv håndtering av avgrensning.",
                "Nivå 5: Avgrensning integrert i realiseringsprosessen."
            ]
        },
        {
            "id": 6,
            "title": "Nullpunkter og estimater",
            "question": "Hvordan valideres og justeres nullpunkter og estimater under realiseringen?",
            "scale": [
                "Nivå 1: Nullpunkter og estimater valideres ikke.",
                "Nivå 2: Store avvik håndteres reaktivt.",
                "Nivå 3: Systematisk validering under realisering.",
                "Nivå 4: Kontinuerlig justering basert på realisering.",
                "Nivå 5: Dynamisk oppdatering av nullpunkter og estimater."
            ]
        },
        {
            "id": 7,
            "title": "Hypotesetesting og datagrunnlag",
            "question": "Hvordan valideres hypoteser og datagrunnlag under realiseringen?",
            "scale": [
                "Nivå 1: Hypoteser valideres ikke under realisering.",
                "Nivå 2: Noen hypoteser testes uformelt.",
                "Nivå 3: Systematisk testing av kritiske hypoteser.",
                "Nivå 4: Omfattende validering under realisering.",
                "Nivå 5: Kontinuerlig hypotesetesting og læring."
            ]
        },
        {
            "id": 8,
            "title": "Interessentengasjement",
            "question": "Hvordan opprettholdes interessentengasjement under realiseringen?",
            "scale": [
                "Nivå 1: Interessentengasjement avtar under realisering.",
                "Nivå 2: Begrenset engasjement for viktige beslutninger.",
                "Nivå 3: Regelmessig dialog med interessenter.",
                "Nivå 4: Aktivt interessentengasjement gjennom realisering.",
                "Nivå 5: Interessenter er drivkrefter i realiseringsarbeidet."
            ]
        },
        {
            "id": 9,
            "title": "Gevinstforutsetninger",
            "question": "Hvordan overvåkes og realiseres gevinstforutsetninger under realiseringen?",
            "scale": [
                "Nivå 1: Forutsetninger overvåkes ikke under realisering.",
                "Nivå 2: Noen forutsetninger følges opp.",
                "Nivå 3: Systematisk overvåkning av forutsetninger.",
                "Nivå 4: Aktiv realisering av forutsetninger.",
                "Nivå 5: Forutsetningsrealisering integrert i gevinstuttak."
            ]
        },
        {
            "id": 10,
            "title": "Prinsipielle og vilkårsmessige kriterier",
            "question": "Hvordan håndteres prinsipielle og vilkårsmessige kriterier under realiseringen?",
            "scale": [
                "Nivå 1: Kriterier håndteres ikke under realisering.",
                "Nivå 2: Store avvik fra kriterier håndteres.",
                "Nivå 3: Systematisk håndtering av kriterier.",
                "Nivå 4: Proaktiv tilpasning til kriterier.",
                "Nivå 5: Kriterier integrert i realiseringsbeslutninger."
            ]
        },
        {
            "id": 11,
            "title": "Enighet om nullpunkter/estimater",
            "question": "Hvordan opprettholdes enighet om nullpunkter og estimater under realiseringen?",
            "scale": [
                "Nivå 1: Enighet testes ikke under realisering.",
                "Nivå 2: Enighet bekreftes ved store endringer.",
                "Nivå 3: Regelmessig bekreftelse av enighet.",
                "Nivå 4: Kontinuerlig arbeid for å opprettholde enighet.",
                "Nivå 5: Full enighet gjennom hele realiseringsfasen."
            ]
        },
        {
            "id": 12,
            "title": "Disponering av kostnads- og tidsbesparelser",
            "question": "Hvordan håndteres disponering av besparelser under realiseringen?",
            "scale": [
                "Nivå 1: Disponering håndteres ikke under realisering.",
                "Nivå 2: Store endringer i disponering håndteres.",
                "Nivå 3: Systematisk revisjon av disponering.",
                "Nivå 4: Dynamisk tilpasning av disponering.",
                "Nivå 5: Optimal disponering under realisering."
            ]
        },
        {
            "id": 13,
            "title": "Måling av effektivitet og produktivitet",
            "question": "Hvordan måles og forbedres effektivitet og produktivitet under realiseringen?",
            "scale": [
                "Nivå 1: Effektivitet og produktivitet måles ikke.",
                "Nivå 2: Noen målinger registreres.",
                "Nivå 3: Systematisk måling og rapportering.",
                "Nivå 4: Målinger brukes til forbedring.",
                "Nivå 5: Kontinuerlig forbedring basert på målinger."
            ]
        },
        {
            "id": 14,
            "title": "Operasjonell risiko og ulemper",
            "question": "Hvordan håndteres operasjonelle risikoer og ulemper under realiseringen?",
            "scale": [
                "Nivå 1: Risikoer og ulemper håndteres ikke.",
                "Nivå 2: Store risikoer håndteres reaktivt.",
                "Nivå 3: Systematisk identifisering og håndtering.",
                "Nivå 4: Proaktiv risikohåndtering.",
                "Nivå 5: Risikostyring integrert i realiseringsarbeid."
            ]
        },
        {
            "id": 15,
            "title": "Balanse mellom gevinster og ulemper",
            "question": "Hvordan vurderes balansen mellom gevinster og ulemper under realiseringen?",
            "scale": [
                "Nivå 1: Balansen vurderes ikke under realisering.",
                "Nivå 2: Balansen vurderes ved store endringer.",
                "Nivå 3: Regelmessig vurdering av balansen.",
                "Nivå 4: Systematisk overvåkning av balansen.",
                "Nivå 5: Balansevurdering integrert i beslutninger."
            ]
        },
        {
            "id": 16,
            "title": "Dokumentasjon og gevinstrealiseringsplan",
            "question": "Hvordan brukes gevinstrealiseringsplanen under realiseringen?",
            "scale": [
                "Nivå 1: Gevinstrealiseringsplanen brukes ikke.",
                "Nivå 2: Planen refereres til ved behov.",
                "Nivå 3: Planen brukes aktivt i realisering.",
                "Nivå 4: Planen oppdateres og brukes kontinuerlig.",
                "Nivå 5: Planen er sentralt styringsverktøy."
            ]
        },
        {
            "id": 17,
            "title": "Gevinstrealiseringsplan som operativ handlingsplan",
            "question": "Hvordan fungerer gevinstrealiseringsplanen som operativ handlingsplan under realiseringen?",
            "scale": [
                "Nivå 1: Planen brukes ikke som operativ handlingsplan.",
                "Nivå 2: Planen brukes til enkelte operasjoner.",
                "Nivå 3: Planen er integrert i operativ styring.",
                "Nivå 4: Planen er aktivt operativt verktøy.",
                "Nivå 5: Planen driver operativ virksomhet."
            ]
        },
        {
            "id": 18,
            "title": "Endringsberedskap og operativ mottaksevne",
            "question": "Hvordan utvikles endringsberedskap og mottaksevne under realiseringen?",
            "scale": [
                "Nivå 1: Endringsberedskap utvikles ikke.",
                "Nivå 2: Begrenset fokus på endringsberedskap.",
                "Nivå 3: Systematisk arbeid med endringsberedskap.",
                "Nivå 4: Målrettet utvikling av mottaksevne.",
                "Nivå 5: Høy mottaksevne og endringsberedskap."
            ]
        },
        {
            "id": 19,
            "title": "Kommunikasjon og forankring",
            "question": "Hvordan opprettholdes kommunikasjon og forankring under realiseringen?",
            "scale": [
                "Nivå 1: Kommunikasjon avtar under realisering.",
                "Nivå 2: Begrenset kommunikasjon om realisering.",
                "Nivå 3: Regelmessig kommunikasjon om fremdrift.",
                "Nivå 4: Systematisk kommunikasjon om realisering.",
                "Nivå 5: Kontinuerlig dialog om realiseringsarbeid."
            ]
        },
        {
            "id": 20,
            "title": "Eierskap og ansvar",
            "question": "Hvordan utøves eierskap og ansvar under realiseringen?",
            "scale": [
                "Nivå 1: Eierskap og ansvar svekkes.",
                "Nivå 2: Begrenset eierskap i realiseringsfasen.",
                "Nivå 3: Tydelig eierskap for realisering.",
                "Nivå 4: Aktivt utøvd eierskap.",
                "Nivå 5: Sterk eierskapskultur i realisering."
            ]
        },
        {
            "id": 21,
            "title": "Periodisering og forankring",
            "question": "Hvordan justeres periodisering og forankring under realiseringen?",
            "scale": [
                "Nivå 1: Periodisering justeres ikke.",
                "Nivå 2: Store justeringer i periodisering.",
                "Nivå 3: Regelmessig revisjon av periodisering.",
                "Nivå 4: Dynamisk tilpasning av periodisering.",
                "Nivå 5: Fleksibel periodisering under realisering."
            ]
        },
        {
            "id": 22,
            "title": "Realisme og engasjement",
            "question": "Hvordan opprettholdes realisme og engasjement under realiseringen?",
            "scale": [
                "Nivå 1: Realisme og engasjement avtar.",
                "Nivå 2: Begrenset fokus på å opprettholde engasjement.",
                "Nivå 3: Arbeid med å opprettholde realisme og engasjement.",
                "Nivå 4: Systematisk arbeid for å styrke troverdighet.",
                "Nivå 5: Høy troverdighet og engasjement."
            ]
        },
        {
            "id": 23,
            "title": "Bygge momentum og tidlig gevinstuttak",
            "question": "Hvordan brukes tidlig gevinstuttak for å bygge momentum i realiseringsfasen?",
            "scale": [
                "Nivå 1: Ingen systematisk bruk av tidlig gevinstuttak.",
                "Nivå 2: Enkelte suksesser brukes til å motivere.",
                "Nivå 3: Bevissthet på viktigheten av momentum.",
                "Nivå 4: Strategisk bruk av tidlige gevinster.",
                "Nivå 5: Momentum systematisk bygget og vedlikeholdt."
            ]
        }
    ],
    "Realisert": [
        {
            "id": 1,
            "title": "Bruk av tidligere læring og gevinstdata",
            "question": "Hvordan dokumenteres og deles læring fra gevinstrealiseringen for fremtidig bruk?",
            "scale": [
                "Nivå 1: Ingen dokumentasjon eller deling av læring.",
                "Nivå 2: Enkelte erfaringer deles uformelt.",
                "Nivå 3: Systematisk dokumentasjon av læring.",
                "Nivå 4: Læring deles og diskuteres i organisasjonen.",
                "Nivå 5: Læring integrert i organisasjonens kunnskapsbase."
            ]
        },
        {
            "id": 2,
            "title": "Strategisk retning og gevinstforståelse",
            "question": "Hvordan bidro den strategiske retningen til gevinstrealiseringens suksess?",
            "scale": [
                "Nivå 1: Strategisk retning bidro lite til suksess.",
                "Nivå 2: Strategi var viktig for enkelte gevinster.",
                "Nivå 3: Strategi bidro til flere gevinster.",
                "Nivå 4: Strategi var avgjørende for gevinstrealisering.",
                "Nivå 5: Strategi og gevinstrealisering fullt integrert."
            ]
        },
        {
            "id": 3,
            "title": "Gevinstkart og visualisering",
            "question": "Hvordan bidro gevinstkartet til gevinstrealiseringens suksess?",
            "scale": [
                "Nivå 1: Gevinstkartet bidro lite til suksess.",
                "Nivå 2: Kartet var nyttig for enkelte gevinster.",
                "Nivå 3: Kartet bidro til flere gevinster.",
                "Nivå 4: Kartet var viktig for gevinstrealisering.",
                "Nivå 5: Kartet var avgjørende for suksess."
            ]
        },
        {
            "id": 4,
            "title": "Strategisk kobling og KPI-er",
            "question": "Hvordan bidro den strategiske koblingen og KPI-ene til gevinstrealisering?",
            "scale": [
                "Nivå 1: Strategisk kobling bidro lite.",
                "Nivå 2: Kobling var viktig for enkelte gevinster.",
                "Nivå 3: Kobling bidro til flere gevinster.",
                "Nivå 4: Kobling var avgjørende for realisering.",
                "Nivå 5: Full integrasjon mellom strategi og realisering."
            ]
        },
        {
            "id": 5,
            "title": "Avgrensning av programgevinst",
            "question": "Hvordan bidro avgrensningsarbeidet til gevinstrealiseringens troverdighet?",
            "scale": [
                "Nivå 1: Avgrensning bidro lite til troverdighet.",
                "Nivå 2: Avgrensning viktig for enkelte gevinster.",
                "Nivå 3: Avgrensning bidro til troverdighet for flere gevinster.",
                "Nivå 4: Avgrensning var avgjørende for troverdighet.",
                "Nivå 5: Avgrensning styrket troverdighet betydelig."
            ]
        },
        {
            "id": 6,
            "title": "Nullpunkter og estimater",
            "question": "Hvordan bidro nullpunkter og estimater til gevinstrealiseringens nøyaktighet?",
            "scale": [
                "Nivå 1: Nullpunkter og estimater bidro lite.",
                "Nivå 2: Estimater var nøyaktige for enkelte gevinster.",
                "Nivå 3: Estimater var nøyaktige for flere gevinster.",
                "Nivå 4: Høy nøyaktighet i estimater.",
                "Nivå 5: Estimater var svært nøyaktige."
            ]
        },
        {
            "id": 7,
            "title": "Hypotesetesting og datagrunnlag",
            "question": "Hvordan bidro hypotesetesting og datagrunnlag til gevinstrealiseringens kvalitet?",
            "scale": [
                "Nivå 1: Testing og datagrunnlag bidro lite.",
                "Nivå 2: Testing viktig for enkelte gevinster.",
                "Nivå 3: Testing bidro til kvalitet for flere gevinster.",
                "Nivå 4: Testing var avgjørende for kvalitet.",
                "Nivå 5: Testing og datagrunnlag styrket kvalitet betydelig."
            ]
        },
        {
            "id": 8,
            "title": "Interessentengasjement",
            "question": "Hvordan bidro interessentengasjement til gevinstrealiseringens suksess?",
            "scale": [
                "Nivå 1: Interessentengasjement bidro lite.",
                "Nivå 2: Engasjement viktig for enkelte gevinster.",
                "Nivå 3: Engasjement bidro til flere gevinster.",
                "Nivå 4: Engasjement var avgjørende for suksess.",
                "Nivå 5: Interessenter var drivkrefter for suksess."
            ]
        },
        {
            "id": 9,
            "title": "Gevinstforutsetninger",
            "question": "Hvordan bidro håndtering av gevinstforutsetninger til realiseringens suksess?",
            "scale": [
                "Nivå 1: Forutsetningshåndtering bidro lite.",
                "Nivå 2: Håndtering viktig for enkelte gevinster.",
                "Nivå 3: Håndtering bidro til flere gevinster.",
                "Nivå 4: Håndtering var avgjørende for suksess.",
                "Nivå 5: Forutsetningshåndtering var suksessfaktor."
            ]
        },
        {
            "id": 10,
            "title": "Prinsipielle og vilkårsmessige kriterier",
            "question": "Hvordan bidro håndtering av kriterier til gevinstrealisering?",
            "scale": [
                "Nivå 1: Kriteriehåndtering bidro lite.",
                "Nivå 2: Håndtering viktig for enkelte gevinster.",
                "Nivå 3: Håndtering bidro til flere gevinster.",
                "Nivå 4: Håndtering var avgjørende for realisering.",
                "Nivå 5: Kriteriehåndtering styrket realisering."
            ]
        },
        {
            "id": 11,
            "title": "Enighet om nullpunkter/estimater",
            "question": "Hvordan bidro enighet om nullpunkter og estimater til realiseringens suksess?",
            "scale": [
                "Nivå 1: Enighet bidro lite til suksess.",
                "Nivå 2: Enighet viktig for enkelte gevinster.",
                "Nivå 3: Enighet bidro til flere gevinster.",
                "Nivå 4: Enighet var avgjørende for suksess.",
                "Nivå 5: Full enighet styrket suksess betydelig."
            ]
        },
        {
            "id": 12,
            "title": "Disponering av kostnads- og tidsbesparelser",
            "question": "Hvordan bidro disponering av besparelser til gevinstrealiseringens verdiskapning?",
            "scale": [
                "Nivå 1: Disponering bidro lite til verdiskapning.",
                "Nivå 2: Disponering viktig for enkelte gevinster.",
                "Nivå 3: Disponering bidro til verdi for flere gevinster.",
                "Nivå 4: Disponering var avgjørende for verdiskapning.",
                "Nivå 5: Optimal disponering maksimerte verdi."
            ]
        },
        {
            "id": 13,
            "title": "Måling av effektivitet og produktivitet",
            "question": "Hvordan bidro måling av effektivitet og produktivitet til gevinstrealisering?",
            "scale": [
                "Nivå 1: Måling bidro lite til realisering.",
                "Nivå 2: Måling viktig for enkelte gevinster.",
                "Nivå 3: Måling bidro til flere gevinster.",
                "Nivå 4: Måling var avgjørende for realisering.",
                "Nivå 5: Måling drevet gevinstrealisering."
            ]
        },
        {
            "id": 14,
            "title": "Operasjonell risiko og ulemper",
            "question": "Hvordan bidro håndtering av risiko og ulemper til gevinstrealiseringens robusthet?",
            "scale": [
                "Nivå 1: Risikohåndtering bidro lite.",
                "Nivå 2: Håndtering viktig for enkelte gevinster.",
                "Nivå 3: Håndtering bidro til robusthet for flere gevinster.",
                "Nivå 4: Håndtering var avgjørende for robusthet.",
                "Nivå 5: Risikohåndtering styrket robusthet betydelig."
            ]
        },
        {
            "id": 15,
            "title": "Balanse mellom gevinster og ulemper",
            "question": "Hvordan bidro balansevurdering til gevinstrealiseringens bærekraft?",
            "scale": [
                "Nivå 1: Balansevurdering bidro lite.",
                "Nivå 2: Vurdering viktig for enkelte gevinster.",
                "Nivå 3: Vurdering bidro til bærekraft for flere gevinster.",
                "Nivå 4: Vurdering var avgjørende for bærekraft.",
                "Nivå 5: Balansevurdering styrket bærekraft betydelig."
            ]
        },
        {
            "id": 16,
            "title": "Dokumentasjon og gevinstrealiseringsplan",
            "question": "Hvordan bidro gevinstrealiseringsplanen til gevinstrealiseringens suksess?",
            "scale": [
                "Nivå 1: Planen bidro lite til suksess.",
                "Nivå 2: Planen viktig for enkelte gevinster.",
                "Nivå 3: Planen bidro til flere gevinster.",
                "Nivå 4: Planen var avgjørende for suksess.",
                "Nivå 5: Planen var suksessfaktor for realisering."
            ]
        },
        {
            "id": 17,
            "title": "Gevinstrealiseringsplan som operativ handlingsplan",
            "question": "Hvordan bidro gevinstrealiseringsplanen som operativ handlingsplan til suksess?",
            "scale": [
                "Nivå 1: Planen som handlingsplan bidro lite.",
                "Nivå 2: Planen viktig for enkelte operasjoner.",
                "Nivå 3: Planen bidro til flere operasjoner.",
                "Nivå 4: Planen var avgjørende for operativ suksess.",
                "Nivå 5: Planen drevet operativ gevinstrealisering."
            ]
        },
        {
            "id": 18,
            "title": "Endringsberedskap og operativ mottaksevne",
            "question": "Hvordan bidro endringsberedskap og mottaksevne til gevinstrealisering?",
            "scale": [
                "Nivå 1: Beredskap og mottaksevne bidro lite.",
                "Nivå 2: Beredskap viktig for enkelte gevinster.",
                "Nivå 3: Beredskap bidro til flere gevinster.",
                "Nivå 4: Beredskap var avgjørende for realisering.",
                "Nivå 5: Høy mottaksevne drevet realisering."
            ]
        },
        {
            "id": 19,
            "title": "Kommunikasjon og forankring",
            "question": "Hvordan bidro kommunikasjon og forankring til gevinstrealiseringens suksess?",
            "scale": [
                "Nivå 1: Kommunikasjon bidro lite til suksess.",
                "Nivå 2: Kommunikasjon viktig for enkelte gevinster.",
                "Nivå 3: Kommunikasjon bidro til flere gevinster.",
                "Nivå 4: Kommunikasjon var avgjørende for suksess.",
                "Nivå 5: God kommunikasjon styrket suksess betydelig."
            ]
        },
        {
            "id": 20,
            "title": "Eierskap og ansvar",
            "question": "Hvordan bidro eierskap og ansvar til gevinstrealiseringens suksess?",
            "scale": [
                "Nivå 1: Eierskap og ansvar bidro lite.",
                "Nivå 2: Eierskap viktig for enkelte gevinster.",
                "Nivå 3: Eierskap bidro til flere gevinster.",
                "Nivå 4: Eierskap var avgjørende for suksess.",
                "Nivå 5: Sterkt eierskap drevet suksess."
            ]
        },
        {
            "id": 21,
            "title": "Periodisering og forankring",
            "question": "Hvordan bidro periodisering og forankring til gevinstrealiseringens effektivitet?",
            "scale": [
                "Nivå 1: Periodisering bidro lite til effektivitet.",
                "Nivå 2: Periodisering viktig for enkelte gevinster.",
                "Nivå 3: Periodisering bidro til effektivitet for flere gevinster.",
                "Nivå 4: Periodisering var avgjørende for effektivitet.",
                "Nivå 5: God periodisering maksimerte effektivitet."
            ]
        },
        {
            "id": 22,
            "title": "Realisme og engasjement",
            "question": "Hvordan bidro realisme og engasjement til gevinstrealiseringens troverdighet?",
            "scale": [
                "Nivå 1: Realisme og engasjement bidro lite.",
                "Nivå 2: Realisme viktig for enkelte gevinster.",
                "Nivå 3: Realisme bidro til troverdighet for flere gevinster.",
                "Nivå 4: Realisme var avgjørende for troverdighet.",
                "Nivå 5: Høy troverdighet styrket realisering."
            ]
        },
        {
            "id": 23,
            "title": "Bygge momentum og tidlig gevinstuttak",
            "question": "Hvordan bidro arbeid med momentum og tidlig gevinstuttak til langsiktig suksess?",
            "scale": [
                "Nivå 1: Momentum og tidlig uttak bidro lite.",
                "Nivå 2: Tidlig uttak viktig for enkelte gevinster.",
                "Nivå 3: Tidlig uttak bidro til momentum for flere gevinster.",
                "Nivå 4: Momentum var avgjørende for suksess.",
                "Nivå 5: Momentum og tidlig uttak drevet langsiktig suksess."
            ]
        }
    ]
}

# ============================================================================
# DATALAGRING
# ============================================================================
def load_data():
    if os.path.exists(DATA_FILE):
        try:
            with open(DATA_FILE, 'rb') as f:
                data = pickle.load(f)
                if 'projects' in data and 'initiatives' not in data:
                    data['initiatives'] = data['projects']
                    del data['projects']
                if 'initiatives' not in data:
                    data['initiatives'] = {}
                return data
        except:
            pass
    return {'initiatives': {}}

def save_data(data):
    with open(DATA_FILE, 'wb') as f:
        pickle.dump(data, f)

def get_data():
    if 'app_data' not in st.session_state:
        st.session_state.app_data = load_data()
    if 'initiatives' not in st.session_state.app_data:
        st.session_state.app_data['initiatives'] = {}
    return st.session_state.app_data

def persist_data():
    save_data(st.session_state.app_data)

# ============================================================================
# STYLING
# ============================================================================
st.markdown(f"""
<style>
@import url('https://fonts.googleapis.com/css2?family=Source+Sans+Pro:wght@400;600;700&display=swap');
html, body, [class*="css"] {{ font-family: 'Source Sans Pro', sans-serif; }}

.main-header {{ font-size: 2rem; color: {COLORS['primary_dark']}; text-align: center; margin-bottom: 0.3rem; font-weight: 700; }}
.sub-header {{ font-size: 0.95rem; color: {COLORS['primary']}; text-align: center; margin-bottom: 1.5rem; }}

.info-box {{ background: linear-gradient(135deg, #C4EFFF 0%, {COLORS['gray_light']} 100%); padding: 1rem; border-radius: 8px; border-left: 4px solid {COLORS['primary_light']}; margin: 0.8rem 0; }}
.success-box {{ background: linear-gradient(135deg, #DDFAE2 0%, {COLORS['gray_light']} 100%); padding: 1rem; border-radius: 8px; border-left: 4px solid {COLORS['success']}; margin: 0.8rem 0; }}
.warning-box {{ background: linear-gradient(135deg, rgba(255, 160, 64, 0.15) 0%, {COLORS['gray_light']} 100%); padding: 1rem; border-radius: 8px; border-left: 4px solid {COLORS['warning']}; margin: 0.8rem 0; }}

.metric-card {{ background: {COLORS['gray_light']}; padding: 1rem; border-radius: 8px; border-left: 4px solid {COLORS['primary']}; text-align: center; margin: 0.3rem 0; }}
.metric-value {{ font-size: 1.6rem; font-weight: 700; color: {COLORS['primary_dark']}; }}
.metric-label {{ font-size: 0.75rem; color: #666; text-transform: uppercase; }}

.strength-card {{ background: linear-gradient(135deg, #DDFAE2 0%, {COLORS['gray_light']} 100%); padding: 0.8rem; border-radius: 8px; border-left: 4px solid {COLORS['success']}; margin: 0.4rem 0; }}
.improvement-card {{ background: linear-gradient(135deg, rgba(255, 107, 107, 0.15) 0%, {COLORS['gray_light']} 100%); padding: 0.8rem; border-radius: 8px; border-left: 4px solid {COLORS['danger']}; margin: 0.4rem 0; }}

.stButton > button {{ background: linear-gradient(135deg, {COLORS['primary']} 0%, {COLORS['primary_dark']} 100%); color: white; border: none; border-radius: 6px; padding: 0.5rem 1rem; font-weight: 600; }}
.stProgress > div > div > div > div {{ background: linear-gradient(90deg, {COLORS['primary_light']} 0%, {COLORS['success']} 100%); }}

#MainMenu {{visibility: hidden;}}
footer {{visibility: hidden;}}
</style>
""", unsafe_allow_html=True)

# ============================================================================
# HJELPEFUNKSJONER
# ============================================================================
def get_score_color(score):
    if score >= 4.5: return COLORS['success']
    elif score >= 3.5: return COLORS['primary_light']
    elif score >= 2.5: return COLORS['warning']
    else: return COLORS['danger']

def get_score_text(score):
    if score >= 4.5: return "Høy modenhet"
    elif score >= 3.5: return "God modenhet"
    elif score >= 2.5: return "Moderat modenhet"
    elif score >= 1.5: return "Begrenset modenhet"
    else: return "Lav modenhet"

def get_recommended_questions(mode, selection, phase):
    if mode == "role" and selection in ROLES:
        return ROLES[selection].get('recommended_questions', {}).get(phase, [])
    elif mode == "parameter":
        recommended = set()
        for param_name in selection:
            if param_name in PARAMETERS:
                recommended.update(PARAMETERS[param_name]['questions'])
        return list(recommended)
    return []

def calculate_stats(initiative, benefit_filter=None):
    if not initiative.get('interviews'):
        return None
    
    all_scores = {}
    for phase in PHASES:
        all_scores[phase] = {}
        for q in questions_data[phase]:
            all_scores[phase][q['id']] = []
    
    interview_count = 0
    for interview in initiative['interviews'].values():
        if benefit_filter and benefit_filter != "all":
            if interview.get('info', {}).get('benefit_id') != benefit_filter:
                continue
        interview_count += 1
        
        for phase, questions in interview.get('responses', {}).items():
            for q_id, resp in questions.items():
                if resp.get('score', 0) > 0:
                    all_scores[phase][int(q_id)].append(resp['score'])
    
    stats = {
        'phases': {},
        'questions': {},
        'parameters': {},
        'total_interviews': interview_count,
        'overall_avg': 0,
        'high_maturity': [],
        'low_maturity': []
    }
    
    all_avgs = []
    
    for phase in PHASES:
        phase_scores = []
        stats['questions'][phase] = {}
        
        for q in questions_data[phase]:
            scores = all_scores[phase][q['id']]
            if scores:
                avg = np.mean(scores)
                stats['questions'][phase][q['id']] = {
                    'avg': avg, 'min': min(scores), 'max': max(scores),
                    'count': len(scores), 'title': q['title'], 'question': q['question']
                }
                phase_scores.append(avg)
                all_avgs.append(avg)
                
                item = {'phase': phase, 'question_id': q['id'], 'title': q['title'], 'score': avg}
                if avg >= 4:
                    stats['high_maturity'].append(item)
                elif avg < 3:
                    stats['low_maturity'].append(item)
        
        if phase_scores:
            stats['phases'][phase] = {'avg': np.mean(phase_scores), 'min': min(phase_scores), 'max': max(phase_scores)}
    
    for param_name, param_data in PARAMETERS.items():
        param_scores = []
        for phase in PHASES:
            if phase in stats['questions']:
                for q_id in param_data['questions']:
                    if q_id in stats['questions'][phase]:
                        param_scores.append(stats['questions'][phase][q_id]['avg'])
        if param_scores:
            stats['parameters'][param_name] = {'avg': np.mean(param_scores), 'description': param_data['description']}
    
    if all_avgs:
        stats['overall_avg'] = np.mean(all_avgs)
    
    stats['high_maturity'].sort(key=lambda x: x['score'], reverse=True)
    stats['low_maturity'].sort(key=lambda x: x['score'])
    
    return stats

# ============================================================================
# DIAGRAMMER
# ============================================================================
def create_phase_radar(phase_data):
    if not phase_data:
        return None
    categories = list(phase_data.keys())
    values = [phase_data[c]['avg'] for c in categories]
    if len(categories) < 3:
        fig = go.Figure(data=[go.Bar(x=categories, y=values, marker_color=COLORS['primary'])])
        fig.update_layout(yaxis=dict(range=[0, 5], title="Score"), height=350, margin=dict(l=40, r=40, t=40, b=40))
        return fig
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(r=values + [values[0]], theta=categories + [categories[0]], fill='toself', fillcolor='rgba(0, 83, 166, 0.3)', line=dict(color=COLORS['primary'], width=3)))
    fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 5], tickvals=[1,2,3,4,5])), showlegend=False, height=350, margin=dict(l=80, r=80, t=40, b=40))
    return fig

def create_parameter_radar(param_data):
    if not param_data:
        return None
    categories = list(param_data.keys())
    values = [param_data[c]['avg'] for c in categories]
    if len(categories) < 3:
        fig = go.Figure(data=[go.Bar(x=categories, y=values, marker_color=COLORS['primary_light'])])
        fig.update_layout(yaxis=dict(range=[0, 5], title="Score"), height=350, margin=dict(l=40, r=40, t=40, b=40))
        return fig
    fig = go.Figure()
    fig.add_trace(go.Scatterpolar(r=values + [values[0]], theta=categories + [categories[0]], fill='toself', fillcolor='rgba(100, 200, 250, 0.3)', line=dict(color=COLORS['primary_light'], width=3)))
    fig.update_layout(polar=dict(radialaxis=dict(visible=True, range=[0, 5])), showlegend=False, height=400, margin=dict(l=100, r=100, t=40, b=40))
    return fig

def create_strength_bar_chart(items, max_items=8):
    if not items:
        return None
    items = items[:max_items]
    labels = [f"{item['phase'][:4]}: {item['title'][:25]}..." if len(item['title']) > 25 else f"{item['phase'][:4]}: {item['title']}" for item in items]
    scores = [item['score'] for item in items]
    fig = go.Figure(data=[go.Bar(x=scores, y=labels, orientation='h', marker_color=COLORS['success'], text=[f"{s:.1f}" for s in scores], textposition='outside')])
    fig.update_layout(xaxis=dict(range=[0, 5.5], title="Score"), yaxis=dict(autorange="reversed"), height=max(250, len(items) * 35), margin=dict(l=200, r=50, t=20, b=40))
    return fig

def create_improvement_bar_chart(items, max_items=8):
    if not items:
        return None
    items = items[:max_items]
    labels = [f"{item['phase'][:4]}: {item['title'][:25]}..." if len(item['title']) > 25 else f"{item['phase'][:4]}: {item['title']}" for item in items]
    scores = [item['score'] for item in items]
    fig = go.Figure(data=[go.Bar(x=scores, y=labels, orientation='h', marker_color=COLORS['danger'], text=[f"{s:.1f}" for s in scores], textposition='outside')])
    fig.update_layout(xaxis=dict(range=[0, 5.5], title="Score"), yaxis=dict(autorange="reversed"), height=max(250, len(items) * 35), margin=dict(l=200, r=50, t=20, b=40))
    return fig

# ============================================================================
# RAPPORT-GENERERING - HTML (FUNGERER ALLTID)
# ============================================================================
def generate_html_report(initiative, stats):
    """Generer komplett HTML-rapport - krever ingen eksterne pakker"""
    
    html = f"""<!DOCTYPE html>
<html lang="no">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Modenhetsvurdering - {initiative['name']}</title>
    <style>
        @import url('https://fonts.googleapis.com/css2?family=Source+Sans+Pro:wght@400;600;700&display=swap');
        * {{ margin: 0; padding: 0; box-sizing: border-box; }}
        body {{ font-family: 'Source Sans Pro', sans-serif; padding: 40px; max-width: 1000px; margin: 0 auto; color: #172141; line-height: 1.6; }}
        h1 {{ color: #172141; text-align: center; margin-bottom: 5px; }}
        h2 {{ color: #0053A6; margin: 30px 0 15px; border-bottom: 2px solid #64C8FA; padding-bottom: 5px; }}
        h3 {{ color: #0053A6; margin: 20px 0 10px; }}
        .subtitle {{ text-align: center; color: #0053A6; margin-bottom: 30px; }}
        table {{ width: 100%; border-collapse: collapse; margin: 15px 0; }}
        th {{ background: #0053A6; color: white; padding: 10px; text-align: left; }}
        td {{ padding: 8px 10px; border-bottom: 1px solid #E8E8E8; }}
        tr:nth-child(even) {{ background: #F2FAFD; }}
        .summary-table td:first-child {{ background: #F2FAFD; font-weight: 600; width: 200px; }}
        .strength {{ color: #35DE6D; font-weight: 600; }}
        .improvement {{ color: #FF6B6B; font-weight: 600; }}
        .score-high {{ background: #DDFAE2; padding: 2px 8px; border-radius: 4px; }}
        .score-low {{ background: rgba(255, 107, 107, 0.15); padding: 2px 8px; border-radius: 4px; }}
        .item {{ padding: 8px 12px; margin: 5px 0; border-radius: 6px; }}
        .item-strength {{ background: linear-gradient(135deg, #DDFAE2 0%, #F2FAFD 100%); border-left: 4px solid #35DE6D; }}
        .item-improvement {{ background: linear-gradient(135deg, rgba(255, 107, 107, 0.15) 0%, #F2FAFD 100%); border-left: 4px solid #FF6B6B; }}
        .comment-box {{ background: #F2FAFD; padding: 15px; margin: 10px 0; border-radius: 8px; border-left: 4px solid #64C8FA; }}
        .footer {{ text-align: center; margin-top: 40px; padding-top: 20px; border-top: 1px solid #E8E8E8; color: #666; }}
        @media print {{ body {{ padding: 20px; }} h2 {{ page-break-before: auto; }} }}
    </style>
</head>
<body>
    <h1>Modenhetsvurdering - Gevinstrealisering</h1>
    <p class="subtitle">Bane NOR - Konsern økonomi og digital transformasjon</p>
    
    <h2>1. Sammendrag</h2>
    <table class="summary-table">
        <tr><td>Endringsinitiativ</td><td>{initiative['name']}</td></tr>
        <tr><td>Beskrivelse</td><td>{initiative.get('description', '-')}</td></tr>
        <tr><td>Rapportdato</td><td>{datetime.now().strftime('%d.%m.%Y')}</td></tr>
        <tr><td>Antall intervjuer</td><td>{stats['total_interviews']}</td></tr>
        <tr><td>Samlet modenhet</td><td><strong>{stats['overall_avg']:.2f}</strong> ({get_score_text(stats['overall_avg'])})</td></tr>
    </table>
"""
    
    # Modenhet per fase
    if stats['phases']:
        html += """
    <h2>2. Modenhet per fase</h2>
    <table>
        <tr><th>Fase</th><th>Gjennomsnitt</th><th>Min</th><th>Maks</th></tr>
"""
        for phase, data in stats['phases'].items():
            html += f"        <tr><td>{phase}</td><td><strong>{data['avg']:.2f}</strong></td><td>{data['min']:.2f}</td><td>{data['max']:.2f}</td></tr>\n"
        html += "    </table>\n"
    
    # Styrkeområder
    if stats['high_maturity']:
        html += """
    <h2>3. Styrkeområder (score ≥ 4)</h2>
"""
        for item in stats['high_maturity'][:15]:
            html += f'    <div class="item item-strength"><strong>[{item["phase"]}]</strong> {item["title"]}: <span class="score-high">{item["score"]:.2f}</span></div>\n'
    
    # Forbedringsområder
    if stats['low_maturity']:
        html += """
    <h2>4. Forbedringsområder (score &lt; 3)</h2>
"""
        for item in stats['low_maturity'][:15]:
            html += f'    <div class="item item-improvement"><strong>[{item["phase"]}]</strong> {item["title"]}: <span class="score-low">{item["score"]:.2f}</span></div>\n'
    
    # Parameterresultater
    if stats['parameters']:
        html += """
    <h2>5. Resultater per parameter</h2>
    <table>
        <tr><th>Parameter</th><th>Score</th><th>Beskrivelse</th></tr>
"""
        for name, data in stats['parameters'].items():
            html += f"        <tr><td>{name}</td><td><strong>{data['avg']:.2f}</strong></td><td>{data['description']}</td></tr>\n"
        html += "    </table>\n"
    
    # Detaljerte resultater per spørsmål
    html += """
    <h2>6. Detaljerte resultater per spørsmål</h2>
"""
    for phase in PHASES:
        if phase in stats['questions'] and stats['questions'][phase]:
            html += f"""
    <h3>{phase}</h3>
    <table>
        <tr><th>ID</th><th>Spørsmål</th><th>Score</th><th>Antall</th></tr>
"""
            for q_id, q_data in sorted(stats['questions'][phase].items()):
                html += f"        <tr><td>{q_id}</td><td>{q_data['title']}</td><td><strong>{q_data['avg']:.2f}</strong></td><td>{q_data['count']}</td></tr>\n"
            html += "    </table>\n"
    
    # Intervjuoversikt
    html += """
    <h2>7. Intervjuoversikt</h2>
    <table>
        <tr><th>Dato</th><th>Intervjuobjekt</th><th>Stilling</th><th>Gevinst</th><th>Fase</th><th>Snitt</th></tr>
"""
    for interview in initiative.get('interviews', {}).values():
        info = interview['info']
        total_answered = sum(1 for phase in interview.get('responses', {}).values() for resp in phase.values() if resp.get('score', 0) > 0)
        total_score = sum(resp['score'] for phase in interview.get('responses', {}).values() for resp in phase.values() if resp.get('score', 0) > 0)
        avg = total_score / total_answered if total_answered > 0 else 0
        html += f"        <tr><td>{info.get('date', '')}</td><td>{info.get('interviewee', '')}</td><td>{info.get('role', '')}</td><td>{info.get('benefit_name', 'Generelt')}</td><td>{info.get('phase', '')}</td><td>{avg:.2f if avg > 0 else '-'}</td></tr>\n"
    html += "    </table>\n"
    
    # Kommentarer fra intervjuer
    html += """
    <h2>8. Kommentarer fra intervjuer</h2>
"""
    for iid, interview in initiative.get('interviews', {}).items():
        info = interview['info']
        html += f'    <h3>{info.get("interviewee", "Ukjent")} - {info.get("date", "")}</h3>\n'
        html += f'    <p><em>Gevinst: {info.get("benefit_name", "Generelt")} | Fase: {info.get("phase", "")} | Stilling: {info.get("role", "")}</em></p>\n'
        
        has_comments = False
        for phase, responses in interview.get('responses', {}).items():
            for q_id, resp in responses.items():
                if resp.get('notes', '').strip():
                    has_comments = True
                    q_title = ""
                    for q in questions_data.get(phase, []):
                        if str(q['id']) == str(q_id):
                            q_title = q['title']
                            break
                    html += f'    <div class="comment-box"><strong>[{phase}] {q_title}:</strong><br>{resp["notes"]}</div>\n'
        
        if not has_comments:
            html += '    <p><em>Ingen kommentarer registrert.</em></p>\n'
    
    # Footer
    html += f"""
    <div class="footer">
        Generert {datetime.now().strftime('%d.%m.%Y %H:%M')} | Bane NOR - Konsern økonomi og digital transformasjon
    </div>
</body>
</html>
"""
    return html

# ============================================================================
# RAPPORT-GENERERING - WORD
# ============================================================================
def generate_word_report(initiative, stats):
    """Generer komplett Word-rapport med alle detaljer"""
    if not DOCX_AVAILABLE:
        st.error("python-docx er ikke installert. Kjør: pip install python-docx")
        return None
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        
        doc = Document()
        
        # Tittel
        title = doc.add_heading('Modenhetsvurdering - Gevinstrealisering', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitle = doc.add_paragraph('Bane NOR - Konsern økonomi og digital transformasjon')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        doc.add_paragraph()
        
        # ===== SAMMENDRAG =====
        doc.add_heading('1. Sammendrag', level=1)
        
        summary_table = doc.add_table(rows=5, cols=2)
        summary_table.style = 'Table Grid'
        cells = [
            ('Endringsinitiativ', initiative['name']),
            ('Beskrivelse', initiative.get('description', '-')),
            ('Rapportdato', datetime.now().strftime('%d.%m.%Y')),
            ('Antall intervjuer', str(stats['total_interviews'])),
            ('Samlet modenhet', f"{stats['overall_avg']:.2f} ({get_score_text(stats['overall_avg'])})")
        ]
        for i, (label, value) in enumerate(cells):
            summary_table.rows[i].cells[0].text = label
            summary_table.rows[i].cells[1].text = value
        
        doc.add_paragraph()
        
        # ===== MODENHET PER FASE =====
        if stats['phases']:
            doc.add_heading('2. Modenhet per fase', level=1)
            table = doc.add_table(rows=len(stats['phases'])+1, cols=4)
            table.style = 'Table Grid'
            headers = ['Fase', 'Gjennomsnitt', 'Min', 'Maks']
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
            for i, (phase, data) in enumerate(stats['phases'].items(), 1):
                table.rows[i].cells[0].text = phase
                table.rows[i].cells[1].text = f"{data['avg']:.2f}"
                table.rows[i].cells[2].text = f"{data['min']:.2f}"
                table.rows[i].cells[3].text = f"{data['max']:.2f}"
            doc.add_paragraph()
        
        # ===== STYRKEOMRÅDER =====
        if stats['high_maturity']:
            doc.add_heading('3. Styrkeområder (score >= 4)', level=1)
            for item in stats['high_maturity'][:15]:
                p = doc.add_paragraph()
                p.add_run(f"[{item['phase']}] ").bold = True
                p.add_run(f"{item['title']}: ")
                p.add_run(f"{item['score']:.2f}").bold = True
            doc.add_paragraph()
        
        # ===== FORBEDRINGSOMRÅDER =====
        if stats['low_maturity']:
            doc.add_heading('4. Forbedringsområder (score < 3)', level=1)
            for item in stats['low_maturity'][:15]:
                p = doc.add_paragraph()
                p.add_run(f"[{item['phase']}] ").bold = True
                p.add_run(f"{item['title']}: ")
                p.add_run(f"{item['score']:.2f}").bold = True
            doc.add_paragraph()
        
        # ===== PARAMETERRESULTATER =====
        if stats['parameters']:
            doc.add_heading('5. Resultater per parameter', level=1)
            table = doc.add_table(rows=len(stats['parameters'])+1, cols=3)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = 'Parameter'
            table.rows[0].cells[1].text = 'Score'
            table.rows[0].cells[2].text = 'Beskrivelse'
            for i, (name, data) in enumerate(stats['parameters'].items(), 1):
                table.rows[i].cells[0].text = name
                table.rows[i].cells[1].text = f"{data['avg']:.2f}"
                table.rows[i].cells[2].text = data['description']
            doc.add_paragraph()
        
        # ===== DETALJERTE RESULTATER PER SPØRSMÅL =====
        doc.add_heading('6. Detaljerte resultater per spørsmål', level=1)
        
        for phase in PHASES:
            if phase in stats['questions'] and stats['questions'][phase]:
                doc.add_heading(f'6.{PHASES.index(phase)+1} {phase}', level=2)
                
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                table.rows[0].cells[0].text = 'ID'
                table.rows[0].cells[1].text = 'Spørsmål'
                table.rows[0].cells[2].text = 'Score'
                table.rows[0].cells[3].text = 'Antall'
                
                for q_id, q_data in sorted(stats['questions'][phase].items()):
                    row = table.add_row()
                    row.cells[0].text = str(q_id)
                    row.cells[1].text = q_data['title']
                    row.cells[2].text = f"{q_data['avg']:.2f}"
                    row.cells[3].text = str(q_data['count'])
                
                doc.add_paragraph()
        
        # ===== INTERVJUOVERSIKT =====
        doc.add_heading('7. Intervjuoversikt', level=1)
        
        if initiative.get('interviews'):
            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            headers = ['Dato', 'Intervjuobjekt', 'Stilling', 'Gevinst', 'Fase', 'Snitt']
            for i, header in enumerate(headers):
                table.rows[0].cells[i].text = header
            
            for interview in initiative['interviews'].values():
                info = interview['info']
                total_answered = sum(1 for phase in interview.get('responses', {}).values() for resp in phase.values() if resp.get('score', 0) > 0)
                total_score = sum(resp['score'] for phase in interview.get('responses', {}).values() for resp in phase.values() if resp.get('score', 0) > 0)
                avg = total_score / total_answered if total_answered > 0 else 0
                
                row = table.add_row()
                row.cells[0].text = info.get('date', '')
                row.cells[1].text = info.get('interviewee', '')
                row.cells[2].text = info.get('role', '')
                row.cells[3].text = info.get('benefit_name', 'Generelt')
                row.cells[4].text = info.get('phase', '')
                row.cells[5].text = f"{avg:.2f}" if avg > 0 else '-'
        
        doc.add_paragraph()
        
        # ===== KOMMENTARER FRA INTERVJUER =====
        doc.add_heading('8. Kommentarer fra intervjuer', level=1)
        
        for iid, interview in initiative.get('interviews', {}).items():
            info = interview['info']
            doc.add_heading(f"{info.get('interviewee', 'Ukjent')} - {info.get('date', '')}", level=2)
            doc.add_paragraph(f"Gevinst: {info.get('benefit_name', 'Generelt')} | Fase: {info.get('phase', '')}")
            
            has_comments = False
            for phase, responses in interview.get('responses', {}).items():
                for q_id, resp in responses.items():
                    if resp.get('notes', '').strip():
                        has_comments = True
                        q_title = ""
                        for q in questions_data.get(phase, []):
                            if str(q['id']) == str(q_id):
                                q_title = q['title']
                                break
                        p = doc.add_paragraph()
                        p.add_run(f"[{phase}] {q_title}: ").bold = True
                        p.add_run(resp['notes'])
            
            if not has_comments:
                doc.add_paragraph("Ingen kommentarer registrert.")
            
            doc.add_paragraph()
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph(f"Generert {datetime.now().strftime('%d.%m.%Y %H:%M')} | Bane NOR - Konsern økonomi og digital transformasjon")
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Lagre
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        st.error(f"Feil ved generering av Word-rapport: {str(e)}")
        return None

# ============================================================================
# RAPPORT-GENERERING - PDF
# ============================================================================
def generate_pdf_report(initiative, stats):
    """Generer komplett PDF-rapport"""
    if not REPORTLAB_AVAILABLE:
        st.error("reportlab er ikke installert. Kjør: pip install reportlab")
        return None
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import cm
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
        from reportlab.lib.enums import TA_CENTER, TA_LEFT
        
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=2*cm, leftMargin=2*cm, topMargin=2*cm, bottomMargin=2*cm)
        
        styles = getSampleStyleSheet()
        styles.add(ParagraphStyle(name='CustomTitle', fontSize=20, alignment=TA_CENTER, spaceAfter=10, textColor=colors.HexColor('#172141')))
        styles.add(ParagraphStyle(name='CustomSubtitle', fontSize=12, alignment=TA_CENTER, spaceAfter=20, textColor=colors.HexColor('#0053A6')))
        styles.add(ParagraphStyle(name='SectionHeader', fontSize=14, spaceAfter=10, spaceBefore=15, textColor=colors.HexColor('#172141'), fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='SubHeader', fontSize=12, spaceAfter=8, spaceBefore=10, textColor=colors.HexColor('#0053A6'), fontName='Helvetica-Bold'))
        styles.add(ParagraphStyle(name='CustomBody', fontSize=10, spaceAfter=5))
        styles.add(ParagraphStyle(name='StrengthText', fontSize=10, textColor=colors.HexColor('#35DE6D')))
        styles.add(ParagraphStyle(name='ImprovementText', fontSize=10, textColor=colors.HexColor('#FF6B6B')))
        
        story = []
        
        # Tittel
        story.append(Paragraph('Modenhetsvurdering - Gevinstrealisering', styles['CustomTitle']))
        story.append(Paragraph('Bane NOR - Konsern økonomi og digital transformasjon', styles['CustomSubtitle']))
        story.append(Spacer(1, 20))
        
        # Sammendrag
        story.append(Paragraph('1. Sammendrag', styles['SectionHeader']))
        summary_data = [
            ['Endringsinitiativ', initiative['name']],
            ['Beskrivelse', initiative.get('description', '-')],
            ['Rapportdato', datetime.now().strftime('%d.%m.%Y')],
            ['Antall intervjuer', str(stats['total_interviews'])],
            ['Samlet modenhet', f"{stats['overall_avg']:.2f} ({get_score_text(stats['overall_avg'])})"]
        ]
        t = Table(summary_data, colWidths=[5*cm, 10*cm])
        t.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F2FAFD')),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E8E8E8')),
            ('VALIGN', (0, 0), (-1, -1), 'TOP'),
            ('PADDING', (0, 0), (-1, -1), 8),
        ]))
        story.append(t)
        story.append(Spacer(1, 15))
        
        # Modenhet per fase
        if stats['phases']:
            story.append(Paragraph('2. Modenhet per fase', styles['SectionHeader']))
            phase_data = [['Fase', 'Gjennomsnitt', 'Min', 'Maks']]
            for phase, data in stats['phases'].items():
                phase_data.append([phase, f"{data['avg']:.2f}", f"{data['min']:.2f}", f"{data['max']:.2f}"])
            t = Table(phase_data, colWidths=[5*cm, 3*cm, 3*cm, 3*cm])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0053A6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E8E8E8')),
                ('PADDING', (0, 0), (-1, -1), 6),
            ]))
            story.append(t)
            story.append(Spacer(1, 15))
        
        # Styrkeområder
        if stats['high_maturity']:
            story.append(Paragraph('3. Styrkeområder (score >= 4)', styles['SectionHeader']))
            for item in stats['high_maturity'][:15]:
                story.append(Paragraph(f"<b>[{item['phase']}]</b> {item['title']}: <b>{item['score']:.2f}</b>", styles['StrengthText']))
            story.append(Spacer(1, 15))
        
        # Forbedringsområder
        if stats['low_maturity']:
            story.append(Paragraph('4. Forbedringsområder (score < 3)', styles['SectionHeader']))
            for item in stats['low_maturity'][:15]:
                story.append(Paragraph(f"<b>[{item['phase']}]</b> {item['title']}: <b>{item['score']:.2f}</b>", styles['ImprovementText']))
            story.append(Spacer(1, 15))
        
        # Parameterresultater
        if stats['parameters']:
            story.append(Paragraph('5. Resultater per parameter', styles['SectionHeader']))
            param_data = [['Parameter', 'Score', 'Beskrivelse']]
            for name, data in stats['parameters'].items():
                param_data.append([name, f"{data['avg']:.2f}", data['description'][:50]])
            t = Table(param_data, colWidths=[4*cm, 2*cm, 8*cm])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0053A6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E8E8E8')),
                ('PADDING', (0, 0), (-1, -1), 6),
                ('FONTSIZE', (0, 0), (-1, -1), 9),
            ]))
            story.append(t)
            story.append(Spacer(1, 15))
        
        # Detaljerte resultater
        story.append(PageBreak())
        story.append(Paragraph('6. Detaljerte resultater per spørsmål', styles['SectionHeader']))
        
        for phase in PHASES:
            if phase in stats['questions'] and stats['questions'][phase]:
                story.append(Paragraph(f'{phase}', styles['SubHeader']))
                q_data = [['ID', 'Spørsmål', 'Score', 'Antall']]
                for q_id, qd in sorted(stats['questions'][phase].items()):
                    q_data.append([str(q_id), qd['title'][:40], f"{qd['avg']:.2f}", str(qd['count'])])
                t = Table(q_data, colWidths=[1*cm, 9*cm, 2*cm, 2*cm])
                t.setStyle(TableStyle([
                    ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#64C8FA')),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E8E8E8')),
                    ('PADDING', (0, 0), (-1, -1), 4),
                    ('FONTSIZE', (0, 0), (-1, -1), 8),
                ]))
                story.append(t)
                story.append(Spacer(1, 10))
        
        # Intervjuoversikt
        story.append(PageBreak())
        story.append(Paragraph('7. Intervjuoversikt', styles['SectionHeader']))
        
        if initiative.get('interviews'):
            int_data = [['Dato', 'Intervjuobjekt', 'Gevinst', 'Fase', 'Snitt']]
            for interview in initiative['interviews'].values():
                info = interview['info']
                total_answered = sum(1 for phase in interview.get('responses', {}).values() for resp in phase.values() if resp.get('score', 0) > 0)
                total_score = sum(resp['score'] for phase in interview.get('responses', {}).values() for resp in phase.values() if resp.get('score', 0) > 0)
                avg = total_score / total_answered if total_answered > 0 else 0
                int_data.append([
                    info.get('date', ''),
                    info.get('interviewee', '')[:20],
                    info.get('benefit_name', 'Gen.')[:15],
                    info.get('phase', '')[:10],
                    f"{avg:.2f}" if avg > 0 else '-'
                ])
            t = Table(int_data, colWidths=[2.5*cm, 4*cm, 3.5*cm, 2.5*cm, 2*cm])
            t.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#0053A6')),
                ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
                ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#E8E8E8')),
                ('PADDING', (0, 0), (-1, -1), 5),
                ('FONTSIZE', (0, 0), (-1, -1), 8),
            ]))
            story.append(t)
        
        # Footer
        story.append(Spacer(1, 30))
        story.append(Paragraph(f"Generert {datetime.now().strftime('%d.%m.%Y %H:%M')} | Bane NOR", styles['CustomSubtitle']))
        
        doc.build(story)
        buffer.seek(0)
        return buffer
        
    except Exception as e:
        st.error(f"Feil ved generering av PDF-rapport: {str(e)}")
        return None

# ============================================================================
# HOVEDAPPLIKASJON
# ============================================================================
def main():
    data = get_data()
    
    # Header med tog-ikon
    train_icon = get_icon_svg('train', COLORS['primary'], 40)
    st.markdown(f'''
    <div style="text-align:center;margin-bottom:1.5rem;">
        <div style="display:flex;justify-content:center;align-items:center;gap:15px;">
            {train_icon}
            <h1 style="margin:0;color:{COLORS['primary_dark']};font-size:2rem;font-weight:700;">Modenhetsvurdering</h1>
            {train_icon}
        </div>
        <p style="color:{COLORS['primary']};font-size:0.95rem;margin-top:0.3rem;">Bane NOR - Konsern økonomi og digital transformasjon</p>
    </div>
    ''', unsafe_allow_html=True)
    
    # Navigasjon
    tab1, tab2, tab3, tab4, tab5 = st.tabs(["Om vurderingen", "Endringsinitiativ", "Intervju", "Resultater", "Rapport"])
    
    # TAB 1: OM VURDERINGEN
    with tab1:
        st.markdown(HENSIKT_TEKST)
        
        # Vis faser med ikoner
        st.markdown("---")
        st.markdown("### Faser i gevinstlivssyklusen")
        phase_cols = st.columns(4)
        for i, phase in enumerate(PHASES):
            with phase_cols[i]:
                icon_name = PHASE_ICONS.get(phase, 'planning')
                icon = get_icon_svg(icon_name, COLORS['primary'], 48)
                st.markdown(f'''<div style="text-align:center;padding:15px;background:{COLORS['gray_light']};border-radius:8px;border-top:4px solid {COLORS['primary']};">
                    {icon}
                    <p style="margin:10px 0 0 0;font-weight:600;color:{COLORS['primary_dark']};">{phase}</p>
                </div>''', unsafe_allow_html=True)
        
        st.markdown("---")
        col1, col2 = st.columns(2)
        with col1:
            people_icon = get_icon_svg('people', COLORS['primary'], 28)
            st.markdown(f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:15px;">{people_icon}<span style="font-size:1.2rem;font-weight:700;color:{COLORS["primary_dark"]};">Tilgjengelige roller</span></div>', unsafe_allow_html=True)
            for role_name, role_data in ROLES.items():
                st.markdown(f'''<div style="padding:8px 12px;background:{COLORS['gray_light']};border-radius:6px;margin:5px 0;border-left:3px solid {COLORS['primary_light']};">
                    <strong>{role_name}</strong><br>
                    <small style="color:#666;">{role_data['description']}</small>
                </div>''', unsafe_allow_html=True)
        with col2:
            dashboard_icon = get_icon_svg('dashboard', COLORS['primary'], 28)
            st.markdown(f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:15px;">{dashboard_icon}<span style="font-size:1.2rem;font-weight:700;color:{COLORS["primary_dark"]};">Tilgjengelige parametere</span></div>', unsafe_allow_html=True)
            for param_name, param_data in PARAMETERS.items():
                st.markdown(f'''<div style="padding:8px 12px;background:{COLORS['gray_light']};border-radius:6px;margin:5px 0;border-left:3px solid {COLORS['primary_light']};">
                    <strong>{param_name}</strong><br>
                    <small style="color:#666;">{param_data['description']}</small>
                </div>''', unsafe_allow_html=True)
    
    # TAB 2: ENDRINGSINITIATIV
    with tab2:
        st.markdown("## Endringsinitiativ og gevinster")
        col1, col2 = st.columns([2, 1])
        
        with col2:
            st.markdown("### Nytt endringsinitiativ")
            with st.form("new_initiative"):
                init_name = st.text_input("Navn", placeholder="F.eks. ERTMS Østlandet")
                init_desc = st.text_area("Beskrivelse", height=80)
                if st.form_submit_button("Opprett", use_container_width=True):
                    if init_name:
                        init_id = datetime.now().strftime("%Y%m%d%H%M%S")
                        data['initiatives'][init_id] = {'name': init_name, 'description': init_desc, 'created': datetime.now().isoformat(), 'benefits': {}, 'interviews': {}}
                        persist_data()
                        st.success(f"'{init_name}' opprettet!")
                        st.rerun()
        
        with col1:
            st.markdown("### Mine endringsinitiativ")
            if not data['initiatives']:
                st.info("Ingen endringsinitiativ ennå. Opprett et nytt for å starte.")
            else:
                for init_id, initiative in data['initiatives'].items():
                    num_interviews = len(initiative.get('interviews', {}))
                    num_benefits = len(initiative.get('benefits', {}))
                    with st.expander(f"{initiative['name']} ({num_benefits} gevinster, {num_interviews} intervjuer)"):
                        st.write(f"**Beskrivelse:** {initiative.get('description', 'Ingen')}")
                        st.markdown("#### Gevinster")
                        with st.form(f"add_benefit_{init_id}"):
                            new_benefit = st.text_input("Ny gevinst", key=f"ben_{init_id}")
                            if st.form_submit_button("Legg til"):
                                if new_benefit:
                                    if 'benefits' not in initiative:
                                        initiative['benefits'] = {}
                                    initiative['benefits'][datetime.now().strftime("%Y%m%d%H%M%S%f")] = {'name': new_benefit, 'created': datetime.now().isoformat()}
                                    persist_data()
                                    st.rerun()
                        for ben_id, benefit in initiative.get('benefits', {}).items():
                            col_a, col_b = st.columns([4, 1])
                            col_a.write(f"• {benefit['name']}")
                            if col_b.button("Slett", key=f"del_ben_{ben_id}"):
                                del initiative['benefits'][ben_id]
                                persist_data()
                                st.rerun()
                        st.markdown("---")
                        if st.button("Slett initiativ", key=f"del_{init_id}"):
                            del data['initiatives'][init_id]
                            persist_data()
                            st.rerun()
    
    # TAB 3: INTERVJU
    with tab3:
        st.markdown("## Gjennomfør intervju")
        if not data['initiatives']:
            st.warning("Opprett et endringsinitiativ først")
        else:
            init_options = {p['name']: pid for pid, p in data['initiatives'].items()}
            selected_init_name = st.selectbox("Velg endringsinitiativ", options=list(init_options.keys()))
            selected_init_id = init_options[selected_init_name]
            initiative = data['initiatives'][selected_init_id]
            
            if 'active_interview' not in st.session_state:
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("### Start nytt intervju")
                    benefit_options = {"Generelt for initiativet": "all"}
                    for ben_id, ben in initiative.get('benefits', {}).items():
                        benefit_options[ben['name']] = ben_id
                    selected_benefit_name = st.selectbox("Gevinst", options=list(benefit_options.keys()))
                    selected_benefit_id = benefit_options[selected_benefit_name]
                    selected_phase = st.selectbox("Fase", options=PHASES)
                    
                    # Vis fase-ikon
                    phase_icon_name = PHASE_ICONS.get(selected_phase, 'planning')
                    phase_icon = get_icon_svg(phase_icon_name, COLORS['primary'], 32)
                    st.markdown(f'<div style="display:flex;align-items:center;gap:10px;margin:5px 0 15px 0;">{phase_icon}<span style="color:{COLORS["primary"]};font-weight:600;">{selected_phase}</span></div>', unsafe_allow_html=True)
                    
                    focus_mode = st.radio("Fokusmodus", options=["Rollebasert", "Parameterbasert", "Alle spørsmål"], horizontal=True)
                    
                    selected_role = None
                    selected_params = []
                    recommended = []
                    
                    if focus_mode == "Rollebasert":
                        selected_role = st.selectbox("Rolle", options=list(ROLES.keys()))
                        recommended = get_recommended_questions("role", selected_role, selected_phase)
                        st.caption(ROLES[selected_role]['description'])
                        st.success(f"{len(recommended)} anbefalte spørsmål. Alle 23 tilgjengelige.")
                    elif focus_mode == "Parameterbasert":
                        selected_params = st.multiselect("Parametere", options=list(PARAMETERS.keys()), default=list(PARAMETERS.keys())[:2])
                        recommended = get_recommended_questions("parameter", selected_params, selected_phase)
                        st.success(f"{len(recommended)} anbefalte spørsmål. Alle 23 tilgjengelige.")
                    
                    with st.form("new_interview"):
                        interviewer = st.text_input("Intervjuer")
                        interviewee = st.text_input("Intervjuobjekt")
                        role_title = st.text_input("Stilling")
                        date = st.date_input("Dato", value=datetime.now())
                        if st.form_submit_button("Start intervju", use_container_width=True):
                            if interviewee:
                                interview_id = datetime.now().strftime("%Y%m%d%H%M%S")
                                initiative['interviews'][interview_id] = {
                                    'info': {'interviewer': interviewer, 'interviewee': interviewee, 'role': role_title, 'date': date.strftime('%Y-%m-%d'), 'phase': selected_phase, 'benefit_id': selected_benefit_id, 'benefit_name': selected_benefit_name, 'focus_mode': focus_mode, 'selected_role': selected_role, 'selected_params': selected_params},
                                    'recommended_questions': recommended, 'responses': {}
                                }
                                persist_data()
                                st.session_state['active_interview'] = {'init_id': selected_init_id, 'interview_id': interview_id}
                                st.rerun()
                
                with col2:
                    st.markdown("### Fortsett eksisterende")
                    if initiative['interviews']:
                        interview_options = {f"{i['info']['interviewee']} - {i['info'].get('benefit_name', 'Generelt')} ({i['info'].get('phase', '?')})": iid for iid, i in initiative['interviews'].items()}
                        selected_interview = st.selectbox("Velg intervju", options=list(interview_options.keys()))
                        if st.button("Fortsett", use_container_width=True):
                            st.session_state['active_interview'] = {'init_id': selected_init_id, 'interview_id': interview_options[selected_interview]}
                            st.rerun()
            
            else:
                active = st.session_state['active_interview']
                if active['init_id'] in data['initiatives']:
                    initiative = data['initiatives'][active['init_id']]
                    if active['interview_id'] in initiative['interviews']:
                        interview = initiative['interviews'][active['interview_id']]
                        phase = interview['info'].get('phase', 'Planlegging')
                        recommended = interview.get('recommended_questions', [])
                        
                        st.markdown(f"### Intervju: {interview['info']['interviewee']}")
                        st.caption(f"Gevinst: {interview['info'].get('benefit_name', 'Generelt')} | Fase: {phase} | Stilling: {interview['info']['role']}")
                        
                        if phase not in interview['responses']:
                            interview['responses'][phase] = {}
                        
                        answered = sum(1 for q_id in range(1, 24) if interview['responses'][phase].get(str(q_id), {}).get('score', 0) > 0)
                        st.progress(answered / 23)
                        st.caption(f"Besvart: {answered} av 23")
                        
                        questions = questions_data[phase]
                        recommended_qs = [q for q in questions if q['id'] in recommended]
                        other_qs = [q for q in questions if q['id'] not in recommended]
                        
                        if recommended_qs:
                            st.markdown("### Anbefalte spørsmål")
                            for q in recommended_qs:
                                q_id_str = str(q['id'])
                                if q_id_str not in interview['responses'][phase]:
                                    interview['responses'][phase][q_id_str] = {'score': 0, 'notes': ''}
                                resp = interview['responses'][phase][q_id_str]
                                status = "✓" if resp['score'] > 0 else "○"
                                with st.expander(f"{status} {q['id']}. {q['title']}" + (f" - Nivå {resp['score']}" if resp['score'] > 0 else ""), expanded=(resp['score'] == 0)):
                                    st.markdown(f"**{q['question']}**")
                                    for level in q['scale']:
                                        st.write(f"- {level}")
                                    new_score = st.radio("Nivå:", options=[0,1,2,3,4,5], index=resp['score'], key=f"s_{phase}_{q['id']}", horizontal=True, format_func=lambda x: "Ikke vurdert" if x == 0 else f"Nivå {x}")
                                    new_notes = st.text_area("Notater:", value=resp['notes'], key=f"n_{phase}_{q['id']}", height=80)
                                    if st.button("Lagre", key=f"save_{phase}_{q['id']}"):
                                        interview['responses'][phase][q_id_str] = {'score': new_score, 'notes': new_notes}
                                        persist_data()
                                        st.rerun()
                        
                        if other_qs:
                            st.markdown("### Andre spørsmål")
                            for q in other_qs:
                                q_id_str = str(q['id'])
                                if q_id_str not in interview['responses'][phase]:
                                    interview['responses'][phase][q_id_str] = {'score': 0, 'notes': ''}
                                resp = interview['responses'][phase][q_id_str]
                                status = "✓" if resp['score'] > 0 else "○"
                                with st.expander(f"{status} {q['id']}. {q['title']}" + (f" - Nivå {resp['score']}" if resp['score'] > 0 else ""), expanded=False):
                                    st.markdown(f"**{q['question']}**")
                                    for level in q['scale']:
                                        st.write(f"- {level}")
                                    new_score = st.radio("Nivå:", options=[0,1,2,3,4,5], index=resp['score'], key=f"s_{phase}_{q['id']}", horizontal=True, format_func=lambda x: "Ikke vurdert" if x == 0 else f"Nivå {x}")
                                    new_notes = st.text_area("Notater:", value=resp['notes'], key=f"n_{phase}_{q['id']}", height=80)
                                    if st.button("Lagre", key=f"save_{phase}_{q['id']}"):
                                        interview['responses'][phase][q_id_str] = {'score': new_score, 'notes': new_notes}
                                        persist_data()
                                        st.rerun()
                        
                        col1, col2 = st.columns(2)
                        if col1.button("Avslutt intervju", use_container_width=True):
                            del st.session_state['active_interview']
                            st.rerun()
                        if col2.button("Avbryt", use_container_width=True):
                            del st.session_state['active_interview']
                            st.rerun()
    
    # TAB 4: RESULTATER
    with tab4:
        st.markdown("## Resultater og analyse")
        if not data['initiatives']:
            st.warning("Ingen endringsinitiativ å vise")
        else:
            init_options = {p['name']: pid for pid, p in data['initiatives'].items()}
            selected_init_name = st.selectbox("Velg endringsinitiativ", options=list(init_options.keys()), key="res_init")
            selected_init_id = init_options[selected_init_name]
            initiative = data['initiatives'][selected_init_id]
            
            benefit_filter_options = {"Alle gevinster": "all"}
            for ben_id, ben in initiative.get('benefits', {}).items():
                benefit_filter_options[ben['name']] = ben_id
            benefit_filter_name = st.selectbox("Filtrer på gevinst:", options=list(benefit_filter_options.keys()))
            benefit_filter = benefit_filter_options[benefit_filter_name]
            
            stats = calculate_stats(initiative, benefit_filter if benefit_filter != "all" else None)
            
            if not stats or stats['total_interviews'] == 0:
                st.info("Ingen intervjuer gjennomført ennå")
            else:
                col1, col2, col3, col4 = st.columns(4)
                col1.markdown(f'<div class="metric-card"><div class="metric-label">Intervjuer</div><div class="metric-value">{stats["total_interviews"]}</div></div>', unsafe_allow_html=True)
                col2.markdown(f'<div class="metric-card"><div class="metric-label">Gjennomsnitt</div><div class="metric-value" style="color: {get_score_color(stats["overall_avg"])}">{stats["overall_avg"]:.2f}</div></div>', unsafe_allow_html=True)
                col3.markdown(f'<div class="metric-card" style="border-left-color:{COLORS["success"]}"><div class="metric-label">Styrker</div><div class="metric-value" style="color: {COLORS["success"]}">{len(stats["high_maturity"])}</div></div>', unsafe_allow_html=True)
                col4.markdown(f'<div class="metric-card" style="border-left-color:{COLORS["danger"]}"><div class="metric-label">Forbedring</div><div class="metric-value" style="color: {COLORS["danger"]}">{len(stats["low_maturity"])}</div></div>', unsafe_allow_html=True)
                
                st.markdown("---")
                col1, col2 = st.columns(2)
                with col1:
                    st.markdown("### Modenhet per fase")
                    if stats['phases']:
                        # Vis fase-oversikt med ikoner
                        for phase_name, phase_data in stats['phases'].items():
                            icon_name = PHASE_ICONS.get(phase_name, 'planning')
                            icon = get_icon_svg(icon_name, get_score_color(phase_data['avg']), 24)
                            st.markdown(f'''<div style="display:flex;align-items:center;gap:10px;padding:8px;background:{COLORS['gray_light']};border-radius:6px;margin:4px 0;">
                                {icon}
                                <span style="flex:1;font-weight:600;">{phase_name}</span>
                                <span style="color:{get_score_color(phase_data['avg'])};font-weight:700;">{phase_data['avg']:.2f}</span>
                            </div>''', unsafe_allow_html=True)
                        st.markdown("")
                        fig = create_phase_radar(stats['phases'])
                        if fig:
                            st.plotly_chart(fig, use_container_width=True)
                with col2:
                    st.markdown("### Modenhet per parameter")
                    if stats['parameters']:
                        fig = create_parameter_radar(stats['parameters'])
                        if fig:
                            st.plotly_chart(fig, use_container_width=True)
                
                st.markdown("---")
                col1, col2 = st.columns(2)
                with col1:
                    target_icon = get_icon_svg('target', COLORS['success'], 28)
                    st.markdown(f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:10px;">{target_icon}<span style="font-size:1.3rem;font-weight:700;color:{COLORS["primary_dark"]};">Styrkeområder</span></div>', unsafe_allow_html=True)
                    if stats['high_maturity']:
                        fig = create_strength_bar_chart(stats['high_maturity'])
                        if fig:
                            st.plotly_chart(fig, use_container_width=True)
                        for item in stats['high_maturity'][:5]:
                            phase_icon = get_icon_svg(PHASE_ICONS.get(item['phase'], 'planning'), COLORS['success'], 18)
                            st.markdown(f'<div class="strength-card"><div style="display:flex;align-items:center;gap:8px;">{phase_icon}<span>[{item["phase"]}] {item["title"]}: <strong>{item["score"]:.2f}</strong></span></div></div>', unsafe_allow_html=True)
                    else:
                        st.info("Ingen styrkeområder identifisert")
                with col2:
                    milestone_icon = get_icon_svg('milestones', COLORS['danger'], 28)
                    st.markdown(f'<div style="display:flex;align-items:center;gap:10px;margin-bottom:10px;">{milestone_icon}<span style="font-size:1.3rem;font-weight:700;color:{COLORS["primary_dark"]};">Forbedringsområder</span></div>', unsafe_allow_html=True)
                    if stats['low_maturity']:
                        fig = create_improvement_bar_chart(stats['low_maturity'])
                        if fig:
                            st.plotly_chart(fig, use_container_width=True)
                        for item in stats['low_maturity'][:5]:
                            phase_icon = get_icon_svg(PHASE_ICONS.get(item['phase'], 'planning'), COLORS['danger'], 18)
                            st.markdown(f'<div class="improvement-card"><div style="display:flex;align-items:center;gap:8px;">{phase_icon}<span>[{item["phase"]}] {item["title"]}: <strong>{item["score"]:.2f}</strong></span></div></div>', unsafe_allow_html=True)
                    else:
                        st.success("Ingen kritiske forbedringsområder!")
    
    # TAB 5: RAPPORT
    with tab5:
        st.markdown("## Generer rapport")
        if not data['initiatives']:
            st.warning("Ingen endringsinitiativ")
        else:
            init_options = {p['name']: pid for pid, p in data['initiatives'].items()}
            selected_init_name = st.selectbox("Velg endringsinitiativ", options=list(init_options.keys()), key="rep_init")
            selected_init_id = init_options[selected_init_name]
            initiative = data['initiatives'][selected_init_id]
            stats = calculate_stats(initiative)
            
            if not stats or stats['total_interviews'] == 0:
                st.info("Gjennomfør minst ett intervju først")
            else:
                st.markdown("### Eksportformat")
                st.markdown("Rapportene inkluderer sammendrag, diagrammer, alle spørsmål og kommentarer fra intervjuer.")
                
                col1, col2, col3 = st.columns(3)
                
                with col1:
                    st.markdown("#### CSV")
                    csv_data = []
                    for phase in stats['questions']:
                        for q_id, q_data in stats['questions'][phase].items():
                            csv_data.append({'Fase': phase, 'SpørsmålID': q_id, 'Tittel': q_data['title'], 'Gjennomsnitt': round(q_data['avg'], 2), 'AntallSvar': q_data['count']})
                    csv_df = pd.DataFrame(csv_data)
                    st.download_button("Last ned CSV", data=csv_df.to_csv(index=False, sep=';'), file_name=f"modenhet_{initiative['name']}_{datetime.now().strftime('%Y%m%d')}.csv", mime="text/csv", use_container_width=True)
                
                with col2:
                    st.markdown("#### Word")
                    if DOCX_AVAILABLE:
                        word_buffer = generate_word_report(initiative, stats)
                        if word_buffer:
                            st.download_button("Last ned Word", data=word_buffer, file_name=f"modenhet_{initiative['name']}_{datetime.now().strftime('%Y%m%d')}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", use_container_width=True)
                    else:
                        st.warning("Installerer pakke... Start appen på nytt.")
                        st.code("pip install python-docx", language="bash")
                
                with col3:
                    st.markdown("#### PDF")
                    if REPORTLAB_AVAILABLE:
                        pdf_buffer = generate_pdf_report(initiative, stats)
                        if pdf_buffer:
                            st.download_button("Last ned PDF", data=pdf_buffer, file_name=f"modenhet_{initiative['name']}_{datetime.now().strftime('%Y%m%d')}.pdf", mime="application/pdf", use_container_width=True)
                    else:
                        st.warning("Installerer pakke... Start appen på nytt.")
                        st.code("pip install reportlab", language="bash")
                
                # HTML fallback som alltid fungerer
                st.markdown("---")
                st.markdown("#### HTML-rapport (fungerer alltid)")
                html_report = generate_html_report(initiative, stats)
                st.download_button("Last ned HTML", data=html_report, file_name=f"modenhet_{initiative['name']}_{datetime.now().strftime('%Y%m%d')}.html", mime="text/html", use_container_width=True)
                st.caption("HTML kan åpnes i nettleser og skrives ut som PDF (Ctrl+P)")
                
                st.markdown("---")
                st.markdown("### Intervjuoversikt")
                interview_data = []
                for iid, interview in initiative['interviews'].items():
                    info = interview['info']
                    total_answered = sum(1 for phase in interview.get('responses', {}).values() for resp in phase.values() if resp.get('score', 0) > 0)
                    total_score = sum(resp['score'] for phase in interview.get('responses', {}).values() for resp in phase.values() if resp.get('score', 0) > 0)
                    avg = total_score / total_answered if total_answered > 0 else 0
                    interview_data.append({'Dato': info.get('date', ''), 'Intervjuobjekt': info.get('interviewee', ''), 'Gevinst': info.get('benefit_name', 'Generelt'), 'Fase': info.get('phase', ''), 'Besvarte': total_answered, 'Snitt': round(avg, 2) if avg > 0 else '-'})
                if interview_data:
                    st.dataframe(pd.DataFrame(interview_data), use_container_width=True)
    
    st.markdown("---")
    footer_icon = get_icon_svg('train', COLORS['primary'], 20)
    st.markdown(f'''<div style="text-align:center;color:#666;font-size:0.85rem;padding:10px 0;">
        <div style="display:flex;justify-content:center;align-items:center;gap:8px;">
            {footer_icon}
            <span>Bane NOR - Konsern økonomi og digital transformasjon</span>
            {footer_icon}
        </div>
    </div>''', unsafe_allow_html=True)

if __name__ == "__main__":
    main()
